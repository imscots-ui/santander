#!/usr/bin/env python3
"""
Santander Business Banking — Concept Prototype: Technical Report.

Clean-room document. Original structure authored for this prototype; not
derived from any external template. Factual description of the front-end
concept prototype only.
"""
from docx import Document
from docx.shared import Pt, RGBColor, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

RED   = RGBColor(0xC8, 0x10, 0x2E)
INK   = RGBColor(0x1C, 0x19, 0x17)
STONE = RGBColor(0x57, 0x53, 0x4E)
LIGHT = RGBColor(0x78, 0x71, 0x6C)

doc = Document()

# Base styles
normal = doc.styles['Normal']
normal.font.name = 'Calibri'
normal.font.size = Pt(10.5)
normal.font.color.rgb = INK
normal.paragraph_format.space_after = Pt(6)
normal.paragraph_format.line_spacing = 1.18

def shade(cell, hexfill):
    tcPr = cell._tc.get_or_add_tcPr()
    sh = OxmlElement('w:shd'); sh.set(qn('w:val'), 'clear'); sh.set(qn('w:fill'), hexfill)
    tcPr.append(sh)

def eyebrow(text):
    p = doc.add_paragraph(); p.paragraph_format.space_after = Pt(1)
    r = p.add_run(text.upper()); r.font.size = Pt(8); r.font.bold = True
    r.font.color.rgb = RED; r.font.name = 'Calibri'
    from docx.oxml.ns import qn as _q
    return p

def h1(text, num=None):
    p = doc.add_paragraph(); p.paragraph_format.space_before = Pt(16); p.paragraph_format.space_after = Pt(4)
    r = p.add_run((f"{num}  " if num else "") + text)
    r.font.size = Pt(15); r.font.bold = True; r.font.color.rgb = INK
    # thin rule under heading
    pPr = p._p.get_or_add_pPr(); pbdr = OxmlElement('w:pBdr'); bottom = OxmlElement('w:bottom')
    bottom.set(qn('w:val'),'single'); bottom.set(qn('w:sz'),'6'); bottom.set(qn('w:space'),'4'); bottom.set(qn('w:color'),'E7E5E4')
    pbdr.append(bottom); pPr.append(pbdr)
    return p

def h2(text):
    p = doc.add_paragraph(); p.paragraph_format.space_before = Pt(9); p.paragraph_format.space_after = Pt(2)
    r = p.add_run(text); r.font.size = Pt(11.5); r.font.bold = True; r.font.color.rgb = STONE
    return p

def body(text):
    p = doc.add_paragraph(); r = p.add_run(text); r.font.size = Pt(10.5); r.font.color.rgb = INK
    return p

def bullet(text):
    p = doc.add_paragraph(style='List Bullet'); r = p.add_run(text); r.font.size = Pt(10.5); r.font.color.rgb = INK
    p.paragraph_format.space_after = Pt(2)
    return p

def kv_table(rows, headers, widths):
    t = doc.add_table(rows=1, cols=len(headers)); t.alignment = WD_TABLE_ALIGNMENT.CENTER
    t.style = 'Table Grid'
    hdr = t.rows[0].cells
    for i, htext in enumerate(headers):
        shade(hdr[i], '1C1917')
        pr = hdr[i].paragraphs[0]; run = pr.add_run(htext); run.font.bold = True; run.font.size = Pt(9)
        run.font.color.rgb = RGBColor(0xFF,0xFF,0xFF)
    for row in rows:
        cells = t.add_row().cells
        for i, val in enumerate(row):
            pr = cells[i].paragraphs[0]; run = pr.add_run(val); run.font.size = Pt(9)
            run.font.color.rgb = INK
            if i == 0: run.font.bold = True
    for i, w in enumerate(widths):
        for c in t.columns[i].cells: c.width = Inches(w)
    return t

# ── COVER ────────────────────────────────────────────────────────────────
for _ in range(3): doc.add_paragraph()
p = doc.add_paragraph(); p.alignment = WD_ALIGN_PARAGRAPH.LEFT
r = p.add_run('INTERNAL · CONFIDENTIAL'); r.font.size = Pt(9); r.font.bold = True; r.font.color.rgb = RED
p = doc.add_paragraph(); r = p.add_run('Santander Business Banking'); r.font.size = Pt(28); r.font.bold = True; r.font.color.rgb = INK
p.paragraph_format.space_after = Pt(0)
p = doc.add_paragraph(); r = p.add_run('Concept Prototype — Technical Report'); r.font.size = Pt(18); r.font.color.rgb = STONE
p = doc.add_paragraph(); r = p.add_run('A front-end exploration of paperless workflows for business banking'); r.font.italic = True; r.font.size = Pt(11); r.font.color.rgb = LIGHT
for _ in range(2): doc.add_paragraph()
meta = doc.add_table(rows=4, cols=2)
for (k, v), row in zip([('Author','Alan Davidson'),('Prepared','July 2026'),('Version','1.0'),('Status','Prototype — internal discussion')], meta.rows):
    a = row.cells[0].paragraphs[0]; ra = a.add_run(k); ra.font.bold=True; ra.font.size=Pt(9); ra.font.color.rgb=STONE
    b = row.cells[1].paragraphs[0]; rb = b.add_run(v); rb.font.size=Pt(9); rb.font.color.rgb=INK
meta.columns[0].cells[0].width = Inches(1.4)
doc.add_page_break()

# ── 1 EXEC SUMMARY ─────────────────────────────────────────────────────────
eyebrow('Section 1'); h1('Executive summary', '1')
body('This report describes a self-initiated concept prototype for Santander Business Banking. '
     'The prototype is a front-end only application that demonstrates how paper-based business-banking '
     'processes — mandate changes, bulk payments, account closures, disputes, cross-border onboarding and '
     'more — can be delivered as guided, self-serve, digitally-signed workflows inside the banking app.')
body('It is a demonstration piece for internal discussion and customer research. It has no backend, uses no '
     'real customer data, and makes no live system connections. Its value is as a tangible, clickable '
     'articulation of a paperless service model and the regulatory controls that model must respect.')
kv_table([
    ['Workflows', '17 end-to-end paperless journeys'],
    ['Entity types', '7 (sole trader → limited, LLP, charity, club, society, partnership)'],
    ['Platform', 'React single-file front end, deployed as one self-contained HTML bundle'],
    ['Scope', 'Front-end concept only — no backend, no real data'],
], ['Dimension', 'Summary'], [1.6, 4.6])

# ── 2 BACKGROUND ───────────────────────────────────────────────────────────
eyebrow('Section 2'); h1('Background & problem statement', '2')
body('Many business-banking service events still depend on paper: posted mandate-variation forms, wet-ink '
     'signatures, branch visits to reset access, faxed international payment details, and written requests for '
     'balance certificates that take days to arrive. Each step adds delay, cost and operational risk, and none '
     'of it is convenient for a business customer running day-to-day operations.')
body('The prototype asks a simple question: what would these processes look like if they were re-imagined as '
     'in-app, self-serve, compliant-by-design workflows — with the regulatory checks (identity, authorisation, '
     'confirmation of payee, cooling-off) built into the flow rather than bolted on with paper?')

# ── 3 OBJECTIVES ───────────────────────────────────────────────────────────
eyebrow('Section 3'); h1('Objectives & scope', '3')
h2('Objectives')
for b in ['Demonstrate paperless equivalents for the most paper-heavy business-banking service events.',
          'Show the regulatory controls each workflow must satisfy, expressed inside the user experience.',
          'Provide a clickable artefact for internal review and customer research.',
          'Explore a consistent design language and accessibility standard for the journeys.']:
    bullet(b)
h2('Non-goals (explicitly out of scope)')
for b in ['No backend, database, or integration with live banking systems.',
          'No real customer, account or payment data — all figures are illustrative.',
          'No production security hardening; this is a front-end concept, not a release candidate.',
          'No processing of real payments, identity documents, or regulatory submissions.']:
    bullet(b)

# ── 4 ARCHITECTURE ─────────────────────────────────────────────────────────
eyebrow('Section 4'); h1('System architecture', '4')
body('The prototype is deliberately built as a single React component in one file, compiled by Vite into a '
     'single self-contained HTML bundle. This keeps the whole product readable top-to-bottom and trivial to '
     'deploy or share as one file.')
h2('Key characteristics')
for b in ['Single-file React front end; all state held at the top of one component.',
          'Compiled to one self-contained HTML artefact via Vite (single-file plugin).',
          'Two interchangeable shells — a mobile bottom-nav layout and a desktop sidebar layout — sharing all state.',
          'A workflow-overlay model: each journey is a full-screen, multi-step wizard rendered over the main app.',
          'An entity system that reconfigures accounts, mandate rules, labels and required documents across 7 business types.']:
    bullet(b)
h2('Mandate & authorisation model')
body('Each account carries a signing rule (any-one, any-two, or all signatories). Workflows resolve the '
     'strictest applicable rule across the accounts involved, which determines whether co-signing or a '
     'cooling-off period is required before an instruction executes.')

# ── 5 WORKFLOW CATALOGUE ───────────────────────────────────────────────────
eyebrow('Section 5'); h1('Workflow catalogue', '5')
body('The prototype implements 17 end-to-end workflows. Each replaces a specific paper or branch-based process.')
kv_table([
    ['Account closure', 'Close an account with in-credit and cooling-off checks', 'BCOBS, mandate rules'],
    ['Business details update', 'Change name, address, contact with proof upload', 'KYB'],
    ['Mandate changes', 'Add/remove signatories or change the signing rule', 'SMR, mandate rules'],
    ['Bulk payments / wages', 'Payee book, Confirmation of Payee, scheduled runs', 'CoP, Faster Payments'],
    ['Pre-approved lending', 'Offer, live repayment calculator, draw-down', 'CCA 1974'],
    ['International FX payment', 'Cross-border payment with rate and fee disclosure', 'Cross-border transfer rules'],
    ['Dormant reactivation', 'Reactivate or close a dormant account', 'Dormancy handling'],
    ['Personal/business unlink', 'Separate personal data from a business relationship', 'GDPR'],
    ['Credit ring-fence', 'Exclude personal data from business credit decisioning', 'GDPR Art. 5(1)(c)'],
    ['Signatory ID register', 'KYC/KYB identity register across evidence lists', 'MLR 2017, KYC'],
    ['MTD VAT submission', 'Categorise, review and submit a VAT return', 'HMRC Making Tax Digital'],
    ['Complaint handling', 'Intake, triage, decision and escalation to FOS', 'FCA DISP'],
    ['Standing orders & DDs', 'Set up standing orders; cancel Direct Debits', 'Direct Debit Guarantee'],
    ['Transaction dispute', 'Chargeback / fraud / DD-Guarantee refund routes', 'PSR 2017, scheme rules'],
    ['Intl. beneficiary onboarding', 'Screen and verify a cross-border payee', 'MLR 2017, sanctions/PEP'],
    ['Balance certificate', 'Issue a sealed certificate of balance / reference', 'Evidence document'],
    ['Trusted-device & sessions', 'Review devices, sign-in activity, sign out others', 'PSD2 SCA, SYSC'],
], ['Workflow', 'What it does', 'Regulatory basis'], [1.9, 3.0, 1.5])

# ── 6 COMPLIANCE ───────────────────────────────────────────────────────────
eyebrow('Section 6'); h1('Regulatory & compliance mapping', '6')
body('The workflows are designed around the controls a real implementation would need. The prototype surfaces '
     'these controls in the user experience so reviewers can see where each obligation is met.')
kv_table([
    ['SCA / PSD2 RTS', 'Strong Customer Authentication and step-up before sensitive actions'],
    ['Confirmation of Payee', 'Account-name verification before first payment to a new payee'],
    ['Direct Debit Guarantee', 'Immediate refund route for Direct Debit disputes'],
    ['PSR 2017', 'Unauthorised-transaction refund and chargeback handling'],
    ['FCA BCOBS / Consumer Duty', 'Cooling-off and fair-treatment handling on closures and changes'],
    ['FCA DISP', 'Complaint intake, triage, decision and Financial Ombudsman escalation'],
    ['FCA SYSC 9', 'Actions recorded to an audit trail'],
    ['MLR 2017', 'Sanctions, PEP and adverse-media screening for new payees'],
    ['HMRC MTD', 'Digital VAT return preparation and submission'],
    ['GDPR', 'Data separation, ring-fencing and purpose limitation'],
], ['Framework', 'How the prototype reflects it'], [2.2, 4.0])

# ── 7 ENGINEERING & QUALITY ────────────────────────────────────────────────
eyebrow('Section 7'); h1('Engineering & quality approach', '7')
body('The prototype is developed through a repeatable build pipeline that treats generation, review, gating '
     'and runtime verification as distinct stages:')
for b in ['Scaffold — new workflows are generated from a template that bakes in the file conventions and prevents recurring defects.',
          'Review — a parallel multi-lens review checks architecture, security, engineering, design and state integrity.',
          'Gate — a pre-commit check blocks accessibility, colour-token, markup and build regressions before code leaves.',
          'Runtime verification — the built app is driven headlessly through real journeys, with an automated accessibility (WCAG) audit.']:
    bullet(b)
h2('Design system')
body('A consistent visual language governs spacing, colour, typography and interaction: a defined spacing scale, '
     'the Santander brand red for primary actions, a warm neutral palette, tabular figures for all monetary '
     'values, and a single primary call-to-action per screen.')

# ── 8 ACCESSIBILITY ────────────────────────────────────────────────────────
eyebrow('Section 8'); h1('Accessibility', '8')
body('Accessibility is treated as a build gate, not an afterthought. Interactive controls carry visible focus '
     'states and accessible names; status is never conveyed by colour alone; and the interface is audited against '
     'WCAG 2.1 AA using an automated engine as part of the verification run. The core screens pass the automated '
     'colour-contrast and name/role checks.')

# ── 9 LIMITATIONS ──────────────────────────────────────────────────────────
eyebrow('Section 9'); h1('Limitations', '9')
for b in ['Front-end only: no data is stored, transmitted or processed.',
          'All names, figures, accounts and references are illustrative and fictional.',
          'Regulatory behaviours are represented for demonstration, not certified for production.',
          'The prototype is not security-tested to production standards.']:
    bullet(b)

# ── 10 ROADMAP ─────────────────────────────────────────────────────────────
eyebrow('Section 10'); h1('Indicative next steps', '10')
for b in ['Prioritise the highest-value journeys for a backend-integrated pilot.',
          'Validate the flows with business customers through structured research.',
          'Complete a full accessibility conformance review with assistive-technology testing.',
          'Define the integration, security and data-protection requirements for a production build.']:
    bullet(b)

doc.add_paragraph()
p = doc.add_paragraph(); r = p.add_run('Prototype only — no backend, no real data, no live system connections. Prepared for internal discussion.')
r.font.size = Pt(8); r.font.italic = True; r.font.color.rgb = LIGHT

doc.save('Santander_Technical_Report.docx')
print('WROTE Santander_Technical_Report.docx')
