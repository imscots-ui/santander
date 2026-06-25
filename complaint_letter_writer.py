#!/usr/bin/env python3
"""
Santander Complaint Decision Letter Writer
FCA DISP-compliant · PSR PS23/3-aware

Usage:  python3 complaint_letter_writer.py           (interactive)
        python3 complaint_letter_writer.py --demo     (sample upheld ISA letter)
        python3 complaint_letter_writer.py --demo2    (sample declined fee letter)

Output: [REF]_[SURNAME]_[YYYYMMDD].docx in current directory
"""

import sys, re
from datetime import date, timedelta
from docx import Document
from docx.shared import Pt, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

# ── Brand ──────────────────────────────────────────────────────────────────────
RED   = RGBColor(0xC8, 0x10, 0x2E)
BLACK = RGBColor(0x1A, 0x1A, 0x1A)
GREY  = RGBColor(0x66, 0x66, 0x66)
LGREY = RGBColor(0xAA, 0xAA, 0xAA)

# ── Bank holidays (England & Wales 2025–2027) ──────────────────────────────────
BH = {
    date(2025,  1,  1), date(2025,  4, 18), date(2025,  4, 21),
    date(2025,  5,  5), date(2025,  5, 26), date(2025,  8, 25),
    date(2025, 12, 25), date(2025, 12, 26),
    date(2026,  1,  1), date(2026,  4,  3), date(2026,  4,  6),
    date(2026,  5,  4), date(2026,  5, 25), date(2026,  8, 31),
    date(2026, 12, 25), date(2026, 12, 28),
    date(2027,  1,  1), date(2027,  3, 26), date(2027,  3, 29),
    date(2027,  5,  3), date(2027,  5, 31), date(2027,  8, 30),
    date(2027, 12, 27), date(2027, 12, 28),
}

def biz_add(d, n):
    while n:
        d += timedelta(1)
        if d.weekday() < 5 and d not in BH:
            n -= 1
    return d

def biz_gap(a, b):
    n, d = 0, a + timedelta(1)
    while d <= b:
        if d.weekday() < 5 and d not in BH:
            n += 1
        d += timedelta(1)
    return n

MONTHS = ["January","February","March","April","May","June",
          "July","August","September","October","November","December"]

def fdate(d):
    return f"{d.day} {MONTHS[d.month-1]} {d.year}"

def fmoney(n):
    if n == 0: return "£0.00"
    if n == int(n): return f"£{int(n):,}"
    return f"£{n:,.2f}"

def interest_calc(principal, rate_pct, days):
    return round(principal * rate_pct / 100 / 365 * days, 2)

# ── Terminal UI ────────────────────────────────────────────────────────────────

def hr():  print("\n" + "═" * 64)
def hr2(): print("─" * 64)
def h(t):  hr(); print(f"  {t}"); hr2()

def ask(prompt, default=None):
    hint = f" [{default}]" if default else ""
    v = input(f"  {prompt}{hint}: ").strip()
    return v if v else (default or "")

def pick(prompt, options):
    print(f"\n  {prompt}")
    for i, o in enumerate(options, 1):
        print(f"  {i:>2}.  {o}")
    while True:
        raw = input("  Choice: ").strip()
        if raw.isdigit() and 1 <= int(raw) <= len(options):
            n = int(raw) - 1
            return n, options[n]
        print("  → Invalid — enter the number shown.")

def getdate(prompt, default=None):
    d0 = default or date.today()
    while True:
        v = input(f"  {prompt} (DD/MM/YYYY) [{d0:%d/%m/%Y}]: ").strip()
        if not v: return d0
        try:
            p = v.split("/")
            return date(int(p[2]), int(p[1]), int(p[0]))
        except:
            print("  → Use DD/MM/YYYY  e.g. 10/06/2026")

def getmoney(prompt):
    while True:
        v = input(f"  {prompt} £").strip().replace(",", "")
        try:    return float(v)
        except: print("  → Number please  e.g. 150.00")

def yn(prompt):
    return input(f"  {prompt} (y/n): ").strip().lower().startswith("y")

def multiline(prompt):
    print(f"\n  {prompt}")
    print("  (type your text — blank line when done)")
    lines = []
    while True:
        v = input("  > ").rstrip()
        if not v and lines: break
        if v: lines.append(v)
    return " ".join(lines)

# ── Complaint category definitions ────────────────────────────────────────────

CATS = [
    ("isa",      "ISA Transfer Delay"),
    ("charge",   "Incorrect Charge / Fee"),
    ("payment",  "Payment Error — wrong amount or beneficiary"),
    ("fraud",    "APP Fraud — reimbursement dispute (PSR PS23/3)"),
    ("mandate",  "Mandate Change Error"),
    ("service",  "Poor Service / Communication"),
    ("closure",  "Account Closure Delay"),
    ("other",    "Other / General"),
]
CAT_IDS    = [k for k, _ in CATS]
CAT_LABELS = [l for _, l in CATS]

# What we examined — category-specific bullet points
EXAMINED = {
    "isa": [
        "Your ISA transfer request and the date received by Santander",
        "Our processing records showing each stage of the ISA transfer",
        "The 15 working day industry standard for Cash ISA transfers",
        "Correspondence between Santander and the receiving institution",
        "Interest rates applicable to your ISA during the relevant period",
    ],
    "charge": [
        "Your account statement showing the disputed charge",
        "The product terms and conditions in force when the charge was applied",
        "Santander's published tariff of charges applicable to your account",
        "Internal processing records relating to the basis of the charge",
        "Any prior correspondence between you and Santander about this charge",
    ],
    "payment": [
        "The payment instruction as submitted (sort code, account number, amount)",
        "Faster Payments / CHAPS / Bacs processing records for the transaction",
        "Your account statements for the relevant period",
        "Any system flags or manual intervention at point of processing",
        "All correspondence between you and Santander regarding this payment",
    ],
    "fraud": [
        "The transaction(s) in question and the payment route used",
        "Records of fraud warnings displayed or communicated at the point of payment",
        "Whether the payment was made via Faster Payments or CHAPS (PSR in-scope systems)",
        "Our fraud referral records and intelligence from the receiving bank",
        "The mandatory reimbursement criteria under PSR PS23/3 (effective 7 October 2024)",
    ],
    "mandate": [
        "The account mandate in force at the time of the instruction",
        "Authority documentation provided in connection with the mandate change",
        "Internal processing records for the mandate instruction",
        "Correspondence with all signatories named on the mandate",
        "Santander's mandate change procedures and authorisation requirements",
    ],
    "service": [
        "Your account history and all relevant correspondence for the period",
        "Call recording transcripts and/or branch visit notes (where available)",
        "Internal account notes made by our colleagues",
        "Our service standards and complaint handling procedures",
        "Any prior escalation or complaint raised in connection with this matter",
    ],
    "closure": [
        "Your account closure request and the date it was received",
        "Our account closure procedure and standard processing timescales",
        "Documentation requirements and when they were fulfilled",
        "The full timeline from your closure request to completion",
        "Correspondence sent to you during the closure process",
    ],
    "other": [
        "All account records and correspondence relevant to your complaint",
        "Our internal policies and procedures applicable to your concern",
        "Applicable regulatory and industry standards",
    ],
}

# Findings text — upheld
FINDING_UPHELD = {
    "isa": (
        "Our investigation has established that your ISA transfer was not completed within the "
        "15 working day industry standard applicable to Cash ISA transfers. This delay is "
        "attributable to Santander and resulted in a period during which your funds were "
        "unavailable to earn interest in your new ISA. We accept your complaint in full and "
        "apologise for the inconvenience caused."
    ),
    "charge": (
        "Our investigation has established that the charge applied to your account was not "
        "correctly applied in accordance with the terms and conditions of your product. We "
        "accept that this charge was applied in error and apologise for this."
    ),
    "payment": (
        "Our investigation has established that an error occurred in the processing of this "
        "payment, attributable to Santander. We accept responsibility for the inconvenience "
        "and financial loss this has caused you."
    ),
    "fraud": (
        "Our investigation, conducted under the Payment Systems Regulator's mandatory "
        "reimbursement framework (PS23/3, effective 7 October 2024), has established that "
        "your loss arose from an APP fraud conducted via a PSR in-scope payment system. We "
        "are satisfied that you took reasonable steps to verify the payee before proceeding "
        "and find no evidence of gross negligence on your part. You are therefore entitled "
        "to reimbursement as set out below."
    ),
    "mandate": (
        "Our investigation has established that the mandate change was not processed in "
        "accordance with the authorisation requirements applicable to your account. We accept "
        "that this error is attributable to Santander and apologise for the consequences "
        "this has had."
    ),
    "service": (
        "Our investigation has established that the service provided on this occasion did not "
        "meet the standard you are entitled to expect from Santander. We accept that our "
        "handling of your case caused you unnecessary inconvenience and we apologise "
        "unreservedly."
    ),
    "closure": (
        "Our investigation has established that your account closure request was not processed "
        "within our standard timescale. This delay is attributable to Santander and caused you "
        "inconvenience. We accept your complaint."
    ),
    "other": (
        "Having reviewed all relevant records and correspondence, we have concluded that your "
        "complaint is substantiated. The service you received fell short of what you should "
        "reasonably expect and we apologise for this."
    ),
}

# Findings text — not upheld
FINDING_DENIED = {
    "isa": (
        "Our investigation has established that your ISA transfer was completed within the "
        "15 working day industry standard. Our records show the transfer was initiated promptly "
        "on receipt of your instruction and completed within the required timeframe. We are "
        "therefore unable to uphold your complaint on this occasion."
    ),
    "charge": (
        "Our investigation has established that the charge applied to your account was correctly "
        "applied in accordance with the terms and conditions of your product as notified to you "
        "at account opening and set out in our published tariff of charges. As the charge was "
        "applied correctly, we are unable to refund it on this occasion."
    ),
    "payment": (
        "Our investigation has established that the payment was processed correctly and strictly "
        "in accordance with the payment instruction we received. We have found no error on "
        "Santander's part and we are therefore unable to uphold your complaint."
    ),
    "fraud": (
        "Our investigation under the PSR mandatory reimbursement framework (PS23/3) has "
        "concluded that we are unable to reimburse your loss on this occasion. Our specific "
        "findings are set out below. We acknowledge this will be a difficult outcome and, if "
        "you wish to challenge our decision, you have the right to refer this to the Financial "
        "Ombudsman Service."
    ),
    "mandate": (
        "Our investigation has established that the mandate change was processed in accordance "
        "with the authority documentation provided to Santander at the relevant time. We are "
        "satisfied that our colleagues followed the correct authorisation procedures and are "
        "unable to uphold your complaint."
    ),
    "service": (
        "Our investigation has established that our colleagues acted appropriately and in "
        "accordance with our service standards throughout. While we understand your frustration, "
        "having reviewed all available records we are unable to uphold your complaint on this "
        "occasion."
    ),
    "closure": (
        "Our investigation has established that your account closure was processed within our "
        "standard timescale once all required documentation had been received and verified. "
        "We are unable to uphold your complaint on this occasion."
    ),
    "other": (
        "Having reviewed all relevant records and correspondence, we have concluded that we "
        "are unable to uphold your complaint on this occasion. We are sorry that this is not "
        "the outcome you were hoping for."
    ),
}

# ── Document helpers ──────────────────────────────────────────────────────────

def set_font(run, size=11, bold=False, italic=False, colour=None):
    run.font.name  = "Calibri"
    run.font.size  = Pt(size)
    run.font.bold  = bold
    run.font.italic = italic
    if colour: run.font.color.rgb = colour

def para(doc, text="", size=11, bold=False, italic=False, colour=None,
         align=WD_ALIGN_PARAGRAPH.LEFT, sb=0, sa=6):
    p = doc.add_paragraph()
    p.alignment = align
    p.paragraph_format.space_before = Pt(sb)
    p.paragraph_format.space_after  = Pt(sa)
    if text:
        r = p.add_run(text)
        set_font(r, size, bold, italic, colour)
    return p

def bullet(doc, text, level=0):
    p = doc.add_paragraph(style="List Bullet")
    p.paragraph_format.left_indent = Cm(0.8 + level * 0.5)
    p.paragraph_format.space_after = Pt(3)
    r = p.add_run(text)
    set_font(r, 11)

def rule(doc, colour="C8102E", sz=18):
    p = doc.add_paragraph()
    bdr = OxmlElement("w:pBdr")
    b   = OxmlElement("w:bottom")
    b.set(qn("w:val"),   "single")
    b.set(qn("w:sz"),    str(sz))
    b.set(qn("w:space"), "1")
    b.set(qn("w:color"), colour)
    bdr.append(b)
    p._p.get_or_add_pPr().append(bdr)
    p.paragraph_format.space_before = Pt(0)
    p.paragraph_format.space_after  = Pt(0)

def shade_cell(cell, hex_colour):
    pr = cell._tc.get_or_add_tcPr()
    s  = OxmlElement("w:shd")
    s.set(qn("w:fill"), hex_colour)
    s.set(qn("w:val"),  "clear")
    pr.append(s)

# ── Letter builder ─────────────────────────────────────────────────────────────

def build_letter(d):
    doc = Document()

    for s in doc.sections:
        s.top_margin    = Cm(2.0)
        s.bottom_margin = Cm(2.0)
        s.left_margin   = Cm(2.5)
        s.right_margin  = Cm(2.5)

    cat     = d["cat"]
    dec     = d["decision"]     # upheld | partial | denied
    stage   = d["stage"]        # stage1 | final
    today   = d["letter_date"]
    recv    = d["received"]
    ref     = d["ref"]
    surname = d["last"]
    title   = d["title"]
    handler = d["handler"]
    team    = d["team"]
    cname   = f"{title} {d['first']} {surname}"

    # ── Banner ────────────────────────────────────────────────────────────────
    rule(doc, "C8102E", 40)
    para(doc, "Santander", size=24, bold=True, colour=RED, sb=6, sa=2)
    para(doc, "Business Banking  ·  Complaints", size=10, colour=GREY, sa=6)
    rule(doc, "BBBBBB", 4)

    # ── Sender block ──────────────────────────────────────────────────────────
    para(doc, "Santander UK plc", size=10, colour=GREY, sb=8, sa=1)
    para(doc, "Complaints Team", size=10, colour=GREY, sa=1)
    para(doc, "PO Box 1109", size=10, colour=GREY, sa=1)
    para(doc, "Bradford  BD1 5WG", size=10, colour=GREY, sa=6)
    para(doc, f"Date:       {fdate(today)}", size=10, sa=1)
    para(doc, f"Reference:  {ref}", size=10, sa=1)
    para(doc, f"Account:    ****{d['account_last4']}", size=10, colour=GREY, sa=14)

    # ── Recipient ─────────────────────────────────────────────────────────────
    for line in [cname] + d["address"]:
        if line.strip():
            para(doc, line, size=11, sa=2)
    para(doc, "", sa=12)

    # ── Salutation ────────────────────────────────────────────────────────────
    para(doc, f"Dear {title} {surname},", size=11, bold=True, sa=10)

    # ── Subject line ─────────────────────────────────────────────────────────
    subj = f"Your complaint — Reference {ref}"
    if stage == "final":
        subj += "  —  FINAL RESPONSE"
    para(doc, subj, size=12, bold=True, colour=RED, sa=12)

    # ── Opening paragraph ─────────────────────────────────────────────────────
    days_held = biz_gap(recv, today)
    if stage == "final":
        opening = (
            "This letter is our final response to your complaint. We have reviewed your case "
            "in full, including any points raised since our initial response, and our findings "
            "and decision are set out below."
        )
    else:
        opening = (
            f"Thank you for raising your complaint with us, which we received on {fdate(recv)}. "
            f"We have now completed our investigation and write to set out our findings. "
            f"This response has been issued within {days_held} business day"
            f"{'s' if days_held != 1 else ''} of receipt."
        )
    para(doc, opening, size=11, sa=12)

    # ── Section 1: What you told us ───────────────────────────────────────────
    para(doc, "1.  What you told us", size=12, bold=True, colour=RED, sa=4)
    para(doc, d["customer_summary"], size=11, sa=12)

    # ── Section 2: What we looked at ─────────────────────────────────────────
    para(doc, "2.  What we looked at", size=12, bold=True, colour=RED, sa=4)
    para(doc, "In reaching our decision we reviewed the following:", size=11, sa=4)
    for item in EXAMINED.get(cat, EXAMINED["other"]):
        bullet(doc, item)
    para(doc, "", sa=8)

    # ── Section 3: Our findings ───────────────────────────────────────────────
    para(doc, "3.  Our findings", size=12, bold=True, colour=RED, sa=4)

    if dec == "upheld":
        finding = FINDING_UPHELD.get(cat, FINDING_UPHELD["other"])
    elif dec == "partial":
        finding = "We have partially upheld your complaint. " + FINDING_UPHELD.get(cat, FINDING_UPHELD["other"])
    else:
        finding = FINDING_DENIED.get(cat, FINDING_DENIED["other"])
        if d.get("denial_detail"):
            finding += f"\n\nIn particular: {d['denial_detail']}."

    para(doc, finding, size=11, sa=12)

    # ── Section 4: What we will do / Our decision ─────────────────────────────
    if dec in ("upheld", "partial"):
        para(doc, "4.  What we will do to put this right", size=12, bold=True, colour=RED, sa=4)

        redress   = d.get("redress", 0)
        goodwill  = d.get("goodwill", 0)
        int_amt   = d.get("interest_amount", 0)
        remedies  = []

        if cat == "isa" and redress > 0:
            remedies.append(
                f"We will credit {fmoney(redress)} to your account in compensation for the "
                f"interest foregone during the delayed transfer period."
            )
        elif cat == "charge" and redress > 0:
            remedies.append(f"We will refund {fmoney(redress)} to your account.")
        elif cat == "payment" and redress > 0:
            remedies.append(
                f"We will credit {fmoney(redress)} to your account in respect of the payment error."
            )
        elif cat == "fraud" and redress > 0:
            net = redress - d.get("excess", 0)
            remedies.append(
                f"We will reimburse {fmoney(net)} to your account under the PSR mandatory "
                f"reimbursement rules within 5 business days of the date of this letter."
            )
            if d.get("excess", 0):
                remedies.append(
                    f"A standard excess of {fmoney(d['excess'])} has been applied under "
                    f"PSR PS23/3 guidelines, reducing the reimbursement from {fmoney(redress)} "
                    f"to {fmoney(net)}."
                )
        elif redress > 0:
            remedies.append(f"We will credit {fmoney(redress)} to your account.")

        if int_amt > 0:
            remedies.append(
                f"We will pay interest of {fmoney(int_amt)}, calculated at {d['int_rate']}% "
                f"per annum on {fmoney(d['int_principal'])} over {d['int_days']} days "
                f"(formula: principal × rate ÷ 100 ÷ 365 × days)."
            )

        if goodwill > 0:
            remedies.append(
                f"As a gesture of goodwill, and without any admission of liability beyond the "
                f"above, we will credit a further {fmoney(goodwill)} to your account."
            )

        if d.get("other_remedy"):
            remedies.append(d["other_remedy"])

        if not remedies:
            remedies.append(
                "We will take the corrective action described above and will confirm once "
                "this has been completed."
            )

        for r in remedies:
            bullet(doc, r)

        para(doc, "", sa=4)
        para(doc,
            "Any credit to your account will be processed within 5 business days of the date "
            "of this letter. If you do not see the credit within that timeframe please contact "
            "us immediately quoting the reference number above.",
            size=11, sa=12)

    else:
        # Not upheld — decision box
        para(doc, "4.  Our decision", size=12, bold=True, colour=RED, sa=4)

        tbl = doc.add_table(rows=1, cols=1)
        tbl.style = "Table Grid"
        cell = tbl.rows[0].cells[0]
        shade_cell(cell, "FEF2F2")

        p1 = cell.add_paragraph()
        p1.paragraph_format.space_after = Pt(4)
        r1 = p1.add_run("Decision: complaint not upheld")
        set_font(r1, 11, bold=True, colour=RED)

        p2 = cell.add_paragraph()
        p2.paragraph_format.space_after = Pt(4)
        r2 = p2.add_run(
            "Having considered all information available to us, we are unable to uphold your "
            "complaint. Our reasons are set out in full in the findings section above. We are "
            "genuinely sorry that this is not the outcome you were hoping for."
        )
        set_font(r2, 11)

        para(doc, "", sa=12)

    # ── Section 5: Your rights (FOS) ─────────────────────────────────────────
    if d.get("fos_eligible"):
        para(doc, "5.  Your right to refer to the Financial Ombudsman Service",
             size=12, bold=True, colour=RED, sa=4)

        if stage == "final":
            fos_text = (
                f"As this is our final response, you now have the right to refer your complaint "
                f"to the Financial Ombudsman Service (FOS) if you remain dissatisfied. You must "
                f"do so within six months of the date of this letter ({fdate(today)}). After "
                f"that date the FOS may not be able to consider your complaint."
            )
        else:
            fos_text = (
                "If you remain dissatisfied with our response, you have the right to refer your "
                "complaint to the Financial Ombudsman Service. As we have issued our decision "
                "above, you may refer at any time — you do not need to wait eight weeks."
            )

        para(doc, fos_text, size=11, sa=8)

        tbl = doc.add_table(rows=1, cols=1)
        tbl.style = "Table Grid"
        cell = tbl.rows[0].cells[0]
        shade_cell(cell, "F7F7F7")

        for text, bold_it in [
            ("Financial Ombudsman Service", True),
            ("Exchange Tower, London  E14 9SR", False),
            ("Freephone: 0800 023 4567  (free from mobiles and landlines)", False),
            ("www.financial-ombudsman.org.uk", False),
            ("", False),
            ("The FOS is a free, independent service. Using it will not affect your legal rights.", False),
        ]:
            cp = cell.add_paragraph()
            cp.paragraph_format.space_after = Pt(2)
            r = cp.add_run(text)
            set_font(r, 10, bold=bold_it)

        para(doc, "", sa=12)

    # ── Contact us ────────────────────────────────────────────────────────────
    para(doc, "If you have any questions", size=12, bold=True, colour=RED, sa=4)
    para(doc,
        f"Please contact our Business Banking Complaints team quoting reference {ref}. "
        f"We are available Monday to Friday, 8am–6pm. You can also write to us at the "
        f"address shown above.",
        size=11, sa=14)

    # ── Sign-off ──────────────────────────────────────────────────────────────
    para(doc, "Yours sincerely,", size=11, sa=28)
    para(doc, handler, size=11, bold=True, sa=2)
    para(doc, team, size=11, colour=GREY, sa=2)
    para(doc, "Santander UK plc", size=11, colour=GREY, sa=2)
    para(doc,
        "Santander UK plc is authorised by the Prudential Regulation Authority and regulated "
        "by the Financial Conduct Authority and the Prudential Regulation Authority (FRN: 106054).",
        size=9, colour=LGREY, sa=18)

    rule(doc, "C8102E", 8)

    for line in [
        "Santander UK plc  ·  Registered in England and Wales No. 2294747",
        "Registered Office: 2 Triton Square, Regent's Place, London  NW1 3AN",
        "Calls may be monitored or recorded for security and quality purposes.",
    ]:
        para(doc, line, size=8, colour=LGREY, sb=2, sa=1)

    return doc

# ── Interactive data collection ───────────────────────────────────────────────

def collect():
    d = {}
    print("\n")
    h("SANTANDER COMPLAINT LETTER WRITER  ·  FCA DISP-Compliant")
    print()
    print("  Generates a ready-to-send Word document complaint response letter.")
    print("  Press Enter to accept defaults shown in [brackets].")
    print()

    # 1 — Dates & stage ───────────────────────────────────────────────────────
    h("1 / 8  —  Dates & Stage")
    d["received"]    = getdate("Date complaint received")
    d["letter_date"] = getdate("Letter date", default=date.today())

    d3  = biz_add(d["received"],  3)
    d5  = biz_add(d["received"],  5)
    d56 = biz_add(d["received"], 56)
    print(f"\n  FCA deadlines from {fdate(d['received'])}:")
    print(f"     D+3  summary resolution:  {fdate(d3)}")
    print(f"     D+5  acknowledgement:     {fdate(d5)}")
    print(f"     D+56 final response:      {fdate(d56)}")

    _, sl = pick("Letter stage", [
        "Stage 1 — Initial Response",
        "Final Response (8-week / escalation re-review)",
    ])
    d["stage"] = "stage1" if "Stage 1" in sl else "final"

    # 2 — Complainant ─────────────────────────────────────────────────────────
    h("2 / 8  —  Complainant Details")
    _, d["title"] = pick("Title", ["Mr","Mrs","Ms","Miss","Dr","Mx","Rev"])
    d["first"]         = ask("First name")
    d["last"]          = ask("Last name")
    print("\n  Address — press Enter to skip a line:")
    addr = []
    for label in ["Line 1 (house / street)","Line 2","Town or City","County","Postcode"]:
        v = ask(label)
        if v: addr.append(v)
    d["address"]       = addr
    d["account_last4"] = ask("Last 4 digits of account number", default="XXXX")
    d["ref"]           = ask("Complaint reference", default=f"COMP-{date.today().year}-001")

    # 3 — Category ────────────────────────────────────────────────────────────
    h("3 / 8  —  Complaint Category")
    idx, _ = pick("Select the complaint type", CAT_LABELS)
    d["cat"] = CAT_IDS[idx]

    # 4 — Customer summary ────────────────────────────────────────────────────
    h("4 / 8  —  What the Customer Said  (Section 1 of the letter)")
    d["customer_summary"] = multiline(
        "Summarise the customer's complaint in your own words (anonymised):"
    )

    # 5 — Decision ────────────────────────────────────────────────────────────
    h("5 / 8  —  Decision")
    _, dl = pick("What is the decision?", [
        "Upheld (in full)",
        "Partially Upheld",
        "Not Upheld",
    ])
    if "full" in dl:         d["decision"] = "upheld"
    elif "Partially" in dl:  d["decision"] = "partial"
    else:                    d["decision"] = "denied"

    # 6 — Redress or denial reason ────────────────────────────────────────────
    if d["decision"] in ("upheld", "partial"):
        h("6 / 8  —  Redress")

        if d["cat"] == "fraud":
            d["redress"] = getmoney("Total fraud loss to reimburse")
            d["excess"]  = 100 if yn("Apply PSR £100 standard excess?") else 0
        else:
            d["redress"] = getmoney("Redress / refund amount  (0 if none)")
            d["excess"]  = 0

        if yn("Add interest compensation?"):
            d["int_principal"] = getmoney("Principal amount for interest calculation")
            d["int_rate"]      = float(ask("Annual interest rate %", default="8.0"))
            d["int_days"]      = int(ask("Number of days"))
            d["interest_amount"] = interest_calc(
                d["int_principal"], d["int_rate"], d["int_days"]
            )
            print(f"\n  Interest calculated: {fmoney(d['interest_amount'])}")
        else:
            d["interest_amount"] = 0

        d["goodwill"]    = getmoney("Goodwill payment  (0 if none)")
        d["other_remedy"] = ask("Any additional remedy to include  (leave blank to skip)")
    else:
        h("6 / 8  —  Reason for Not Upholding")
        _, r = pick("Primary reason the complaint is declined:", [
            "Service was provided within the product terms and conditions",
            "No error found on Santander's part",
            "Claim is outside the relevant time limit",
            "Gross negligence — APP fraud (customer did not follow warnings or act reasonably)",
            "Customer complaint is factually incorrect based on our records",
            "Other — I will type the reason below",
        ])
        if "Other" in r or "type" in r:
            r = ask("Type the reason")
        d["denial_detail"] = r
        d["redress"] = d["goodwill"] = d["interest_amount"] = 0

    # 7 — Handler ─────────────────────────────────────────────────────────────
    h("7 / 8  —  Handler Details")
    d["handler"] = ask("Your full name (as it will appear on the letter)")
    d["team"]    = ask("Team / department", default="Business Banking Complaints")

    # 8 — FOS eligibility ─────────────────────────────────────────────────────
    h("8 / 8  —  FOS Eligibility")
    print("  Eligible complainants (DISP 2.7): individuals, micro-enterprises,")
    print("  small charities (income < £1m), small trusts (assets < £1m),")
    print("  small businesses (< 10 staff and ≤ £2m turnover).")
    d["fos_eligible"] = yn("Include FOS referral rights section?")

    # ── Review before generating ──────────────────────────────────────────────
    h("REVIEW — confirm before generating")
    print(f"  Complainant:  {d['title']} {d['first']} {d['last']}")
    print(f"  Reference:    {d['ref']}")
    print(f"  Category:     {dict(CATS)[d['cat']]}")
    print(f"  Decision:     {d['decision'].upper()}")
    print(f"  Stage:        {'Final Response' if d['stage'] == 'final' else 'Stage 1'}")
    if d.get("redress", 0):     print(f"  Redress:      {fmoney(d['redress'])}")
    if d.get("goodwill", 0):    print(f"  Goodwill:     {fmoney(d['goodwill'])}")
    if d.get("interest_amount", 0): print(f"  Interest:     {fmoney(d['interest_amount'])}")
    print(f"  Handler:      {d['handler']}")
    print(f"  FOS section:  {'Yes' if d['fos_eligible'] else 'No'}")
    print()

    if not yn("Generate this letter now?"):
        print("\n  Cancelled.")
        sys.exit(0)

    return d

# ── Demo data sets ─────────────────────────────────────────────────────────────

def demo_upheld():
    return {
        "received":       date(2026, 6, 10),
        "letter_date":    date(2026, 6, 25),
        "stage":          "stage1",
        "title":          "Ms",
        "first":          "Sarah",
        "last":           "Mitchell",
        "address":        ["14 Thornfield Road", "Leeds", "LS6 2AA"],
        "account_last4":  "4471",
        "ref":            "COMP-2026-00142",
        "cat":            "isa",
        "customer_summary": (
            "The customer contacted us to complain that her Cash ISA transfer from a previous "
            "provider took 23 working days to complete. She states the delay caused her to miss "
            "out on interest and that she received no updates during the transfer period. She is "
            "requesting compensation for the lost interest and an apology."
        ),
        "decision":         "upheld",
        "redress":          47.50,
        "excess":           0,
        "interest_amount":  0,
        "goodwill":         25.00,
        "other_remedy":     "",
        "handler":          "James Hartley",
        "team":             "Business Banking Complaints",
        "fos_eligible":     True,
    }

def demo_declined():
    return {
        "received":       date(2026, 6, 10),
        "letter_date":    date(2026, 6, 25),
        "stage":          "stage1",
        "title":          "Mr",
        "first":          "Daniel",
        "last":           "Pearce",
        "address":        ["22 Oak Street", "Manchester", "M4 5LQ"],
        "account_last4":  "8812",
        "ref":            "COMP-2026-00143",
        "cat":            "charge",
        "customer_summary": (
            "The customer has contacted us to dispute a £15.00 charge applied to his Business "
            "Current Account on 10 June 2026. He states he was unaware of this charge and does "
            "not believe it should have been applied to his account. He is requesting a full "
            "refund of £15.00."
        ),
        "decision":       "denied",
        "denial_detail":  (
            "Our records confirm that the £15.00 charge is the standard monthly account fee "
            "for the Business Current Account. This fee is clearly set out in the Product "
            "Summary and Terms and Conditions provided to the customer at account opening on "
            "14 March 2024, and is listed in our published Business Banking tariff. The charge "
            "was applied correctly and we are therefore unable to issue a refund."
        ),
        "redress":         0,
        "goodwill":        0,
        "interest_amount": 0,
        "handler":         "James Hartley",
        "team":            "Business Banking Complaints",
        "fos_eligible":    True,
    }

# ── Main ──────────────────────────────────────────────────────────────────────

def main():
    args = sys.argv[1:]

    if "--demo2" in args:
        print("\n  [DEMO 2] Generating sample declined fee letter...")
        d = demo_declined()
    elif "--demo" in args:
        print("\n  [DEMO 1] Generating sample upheld ISA letter...")
        d = demo_upheld()
    else:
        try:
            d = collect()
        except (KeyboardInterrupt, EOFError):
            print("\n\n  Cancelled.")
            sys.exit(0)

    print("\n  Building letter...")
    doc = build_letter(d)

    safe_ref  = re.sub(r"[^\w\-]", "", d["ref"])
    safe_surn = re.sub(r"\W",      "", d["last"]).upper()
    dt        = d["letter_date"].strftime("%Y%m%d")
    filename  = f"{safe_ref}_{safe_surn}_{dt}.docx"

    doc.save(filename)

    print(f"\n  ✓  Letter saved:  {filename}")
    print(f"     Complainant:   {d['title']} {d['first']} {d['last']}")
    print(f"     Reference:     {d['ref']}")
    print(f"     Category:      {dict(CATS)[d['cat']]}")
    print(f"     Decision:      {d['decision'].upper()}")
    print(f"     Stage:         {'Final Response' if d['stage'] == 'final' else 'Stage 1'}")
    if d.get("redress",         0): print(f"     Redress:       {fmoney(d['redress'])}")
    if d.get("goodwill",        0): print(f"     Goodwill:      {fmoney(d['goodwill'])}")
    if d.get("interest_amount", 0): print(f"     Interest:      {fmoney(d['interest_amount'])}")
    print()


if __name__ == "__main__":
    main()
