#!/usr/bin/env python3
"""
Santander Business Banking Prototype — Position Paper
Mirrors the layout of the Quantum "Beyond Defence" position paper: classification
banner · title block · abstract · Executive summary · Basis and scope of claims
· What it is (+ two-mode table) · The general problem · How it would be used ·
Adjacent segments (obligation -> capability mapping) + summary table · Conclusion
· References. Branded Santander, INTERNAL - CONFIDENTIAL.

Produces: Santander_Position_Paper.docx
"""
from docx import Document
from docx.shared import Pt, RGBColor, Inches, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

RED   = RGBColor(0xC8, 0x10, 0x2E)
DARK  = RGBColor(0x1A, 0x1A, 0x1A)
GREY  = RGBColor(0x66, 0x66, 0x66)
LGREY = RGBColor(0x99, 0x99, 0x99)
WHITE = RGBColor(0xFF, 0xFF, 0xFF)

doc = Document()
for s in doc.sections:
    s.page_width, s.page_height = Inches(8.27), Inches(11.69)
    s.left_margin = s.right_margin = Cm(2.4)
    s.top_margin = s.bottom_margin = Cm(2.2)

base = doc.styles['Normal']
base.font.name = 'Calibri'; base.font.size = Pt(10.5); base.font.color.rgb = DARK
base.paragraph_format.space_after = Pt(6); base.paragraph_format.line_spacing = 1.12

def _shade(el, hexv):
    sh = OxmlElement('w:shd'); sh.set(qn('w:val'), 'clear'); sh.set(qn('w:fill'), hexv); el.append(sh)

def banner(text):
    t = doc.add_table(rows=1, cols=1); t.alignment = WD_TABLE_ALIGNMENT.CENTER
    c = t.rows[0].cells[0]; _shade(c._tc.get_or_add_tcPr(), '1A1A1A')
    p = c.paragraphs[0]; p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p.add_run(text); r.font.bold = True; r.font.size = Pt(9); r.font.color.rgb = WHITE; r.font.name='Calibri'
    p.paragraph_format.space_after = Pt(0); p.paragraph_format.space_before = Pt(0)

def para(text, size=10.5, bold=False, italic=False, color=DARK, align=WD_ALIGN_PARAGRAPH.LEFT, sb=0, sa=6):
    p = doc.add_paragraph(); p.alignment = align
    p.paragraph_format.space_before = Pt(sb); p.paragraph_format.space_after = Pt(sa)
    r = p.add_run(text); r.font.size = Pt(size); r.font.bold = bold; r.font.italic = italic
    r.font.color.rgb = color; r.font.name='Calibri'
    return p

def h1(num, text):
    p = doc.add_paragraph(); p.paragraph_format.space_before = Pt(14); p.paragraph_format.space_after = Pt(4)
    r = p.add_run(f'{num} {text}' if num else text); r.font.size = Pt(14); r.font.bold = True; r.font.color.rgb = DARK; r.font.name='Calibri'
    pPr = p._p.get_or_add_pPr(); bdr = OxmlElement('w:pBdr'); b = OxmlElement('w:bottom')
    b.set(qn('w:val'),'single'); b.set(qn('w:sz'),'6'); b.set(qn('w:space'),'4'); b.set(qn('w:color'),'C8102E')
    bdr.append(b); pPr.append(bdr)

def h2(num, text):
    p = doc.add_paragraph(); p.paragraph_format.space_before = Pt(10); p.paragraph_format.space_after = Pt(3)
    r = p.add_run(f'{num} {text}'); r.font.size = Pt(11.5); r.font.bold = True; r.font.color.rgb = RED; r.font.name='Calibri'

def bullet(text):
    p = doc.add_paragraph(style='List Bullet'); p.paragraph_format.space_after = Pt(3); p.paragraph_format.left_indent = Cm(0.7)
    r = p.add_run(text); r.font.size = Pt(10.5); r.font.color.rgb = DARK; r.font.name='Calibri'

def abstract_box(text):
    t = doc.add_table(rows=1, cols=1); c = t.rows[0].cells[0]; _shade(c._tc.get_or_add_tcPr(), 'F7F2F0')
    tcPr = c._tc.get_or_add_tcPr(); borders = OxmlElement('w:tcBorders')
    lft = OxmlElement('w:left'); lft.set(qn('w:val'),'single'); lft.set(qn('w:sz'),'18'); lft.set(qn('w:space'),'0'); lft.set(qn('w:color'),'C8102E')
    borders.append(lft); tcPr.append(borders)
    p = c.paragraphs[0]
    r = p.add_run('Summary.  '); r.font.bold=True; r.font.size=Pt(10); r.font.name='Calibri'; r.font.color.rgb=DARK
    r2 = p.add_run(text); r2.font.size=Pt(10); r2.font.name='Calibri'; r2.font.color.rgb=DARK
    p.paragraph_format.space_after = Pt(2)

def caption(text):
    p = doc.add_paragraph(); p.paragraph_format.space_before = Pt(3); p.paragraph_format.space_after = Pt(10)
    r = p.add_run(text); r.font.size = Pt(8.5); r.font.italic = True; r.font.color.rgb = GREY; r.font.name='Calibri'

def table(headers, rows, widths=None, capt=None):
    t = doc.add_table(rows=1, cols=len(headers)); t.alignment = WD_TABLE_ALIGNMENT.CENTER; t.style = 'Table Grid'
    hdr = t.rows[0].cells
    for i, htext in enumerate(headers):
        _shade(hdr[i]._tc.get_or_add_tcPr(), 'C8102E')
        p = hdr[i].paragraphs[0]; r = p.add_run(htext); r.font.bold=True; r.font.size=Pt(8.5); r.font.color.rgb=WHITE; r.font.name='Calibri'
        p.paragraph_format.space_after = Pt(1)
    for ri, row in enumerate(rows):
        cells = t.add_row().cells
        for ci, val in enumerate(row):
            if ri % 2 == 0: _shade(cells[ci]._tc.get_or_add_tcPr(), 'F7F7F6')
            p = cells[ci].paragraphs[0]; r = p.add_run(val); r.font.size=Pt(8.5); r.font.color.rgb=DARK; r.font.name='Calibri'
            p.paragraph_format.space_after = Pt(1)
    if widths:
        for ci, w in enumerate(widths):
            for row in t.rows: row.cells[ci].width = Cm(w)
    if capt: caption(capt)
    return t

# ── Banner + title ──────────────────────────────────────────────────────────
banner('INTERNAL · CONFIDENTIAL')
para('Beyond Business Banking: Applications, Adjacent Segments and Cross-Division Relevance of a '
     'Paperless, Compliance-Embedded Workflow Prototype',
     size=18, bold=True, sb=10, sa=4)
para('Alan Davidson', size=11, bold=True, sa=0)
para('Santander Business Banking', size=10, color=GREY, sa=0)
para('Position Paper · June 2026', size=10, italic=True, color=RED, sa=10)

abstract_box(
    'The Business Banking prototype demonstrates paperless equivalents of branch- and post-based '
    'processes, with the governing regulatory control embedded in each workflow and every action written '
    'to an immutable audit trail. This paper argues the approach is not specific to business banking. '
    'Stripped of its business-banking framing, the prototype is a pattern for turning any regulated, '
    'evidence-bearing customer or operational process into a guided, self-serve digital workflow. It '
    'names the three conditions under which the pattern pays off, then maps it — obligation by '
    'obligation — onto adjacent Santander segments and wider financial-services settings. Claims about '
    'the prototype are drawn from its own build and verification record; claims about an obligation refer '
    'only to named, publicly established regulation; and the link between them is presented as a '
    'reasoned structural mapping, not a promise of a finished product.')

para('Keywords: paperless workflows; embedded compliance; audit trail; Consumer Duty; self-serve '
     'banking; process digitisation; SCA; accessibility; cross-segment applicability.',
     size=9, italic=True, color=GREY, sa=8)

# ── 1 Executive summary ──────────────────────────────────────────────────────
h1('1', 'Executive summary')
para('Most regulated customer processes still carry paper, branch visits or manual back-office steps — '
     'and every one of them must be evidenced to a regulator, auditor or the customer themselves. The '
     'Business Banking prototype shows that such a process can instead be a guided digital workflow in '
     'which the control (authorisation, cooling-off, Confirmation of Payee, complaint handling) is part '
     'of the flow, and every step is logged. The value is not the specific workflows; it is the pattern.')
para('This paper (i) states the standard of evidence it holds itself to (§2); (ii) summarises what the '
     'prototype verifiably is and does (§3); (iii) abstracts the general problem it solves into three '
     'conditions (§4); (iv) describes how it would be used (§5); and (v) maps the pattern onto adjacent '
     'Santander segments and the wider sector (§6). It closes with a candid statement of limits (§7).')

# ── 2 Basis and scope of claims ──────────────────────────────────────────────
h1('2', 'Basis and scope of claims')
para('This paper is written to a deliberate evidentiary standard, because its value depends on being '
     'defensible rather than promotional. It contains three kinds of supportable statement:')
bullet('Verified facts about the prototype — drawn from its own build and demonstration record: the '
       'workflows it renders, the controls it embeds, the audit trail it writes, and how it is deployed.')
bullet('Facts about the existence and purpose of named obligations — an FCA sourcebook, a PSD2 technical '
       'standard, a Consumer Duty principle — cited as publicly established regulation.')
bullet('A structural mapping between the two — a reasoned correspondence between a prototype capability '
       'and a real obligation. This is an argument of applicability, not a claim that the prototype is a '
       'certified system for that use.')
para('Where a segment is named in §6, the obligation cited is factual; the mapping onto it is the '
     'structural claim; and nothing here asserts that the prototype, as built, is production-ready for '
     'that segment.')

# ── 3 What the prototype is ──────────────────────────────────────────────────
h1('3', 'What the prototype is')
para('The prototype is a single self-contained HTML file (~890 KB) that runs in any browser with no '
     'installation, no backend and no real data. It presents a business-banking app whose actions are '
     'guided workflows with embedded compliance.')
h2('3.1', 'Three things it does')
bullet('Guides — turns a paper/branch process into a step-based workflow (thirteen today: mandates, bulk '
       'payments, closures, MTD, lending, FX, standing orders & Direct Debits, complaints and more).')
bullet('Controls — embeds the governing regulatory step in the flow: dual authorisation with personal-PIN '
       'signing, working-day cooling-off, Confirmation of Payee, KYC/KYB, DISP complaint handling.')
bullet('Evidences — writes every action to an immutable, actor-and-timestamp audit trail aligned to FCA '
       'SYSC 9, giving the demonstrability the Senior Managers Regime expects.')
h2('3.2', 'Two modes, one control model')
para('The prototype is a demonstration; a production build would keep the same control model behind real '
     'services. Table 1 states the honest boundary between the two.')
table(
    ['Property', 'Demonstration (today)', 'Production path (Phase 2)'],
    [['Data','Fictional, in-memory','Real, from core banking / CRM'],
     ['Back-end','None — front-end only','Node.js services (payments, KYC, HMRC, audit)'],
     ['Controls','Modelled in the flow','Same flow, enforced server-side and assured'],
     ['Audit trail','In-session, illustrative','Append-only store, 7-year retention (SYSC 9)'],
     ['Assurance','Design demonstration','Controlled, tested, independently assured'],
     ['Cost / reach','Free, runs anywhere','Hosted, integrated, access-controlled']],
    widths=[3.2,4.9,5.2],
    capt='Table 1. The prototype demonstrates the control model; production keeps the same workflows and '
         'controls but enforces them behind real, assured services.')
h2('3.3', 'Deployment and rollout')
para('Because it is a static single file, the prototype can be shown on any device, sent to colleagues, '
     'or hosted for customer research. A production rollout would follow the graded path in Table 1, '
     'segment by segment, reusing the workflow and control patterns rather than rebuilding them.')

# ── 4 The general problem ────────────────────────────────────────────────────
h1('4', 'The general problem the prototype solves')
para('Stripped of its business-banking framing, the prototype is an instrument for turning a regulated, '
     'evidence-bearing process into a guided self-serve workflow. It is valuable precisely where three '
     'conditions hold together:')
bullet('A paper, branch or manual process. The task today requires a form, a visit, a phone call or a '
       'back-office hand-off that could instead be self-served.')
bullet('An obligation to evidence a control. A regulator, auditor or the customer expects the '
       'organisation to show that the right check was made, by the right person, at the right time.')
bullet('A driver to remove friction. Cost, speed, accessibility or Consumer-Duty good-outcome pressure '
       'makes the manual process a liability worth removing.')
para('Business banking satisfies all three, which is why the prototype began there. The sections that '
     'follow show that many adjacent Santander segments — and the wider sector — satisfy all three too.')

# ── 5 How it would be used ───────────────────────────────────────────────────
h1('5', 'How it would be used')
para('In practice the pattern is pointed at a process that today carries paper or manual steps, and it '
     'is applied in three moves:')
bullet('Guide the customer or colleague through the task as a step-based workflow, forking by the '
       'circumstances that matter (entity type, signing rule, product, vulnerability).')
bullet('Embed the control at the point it is needed — authorise, verify the payee, start the cooling-off, '
       'check eligibility — with the governing regulation named in the interface.')
bullet('Evidence the outcome by writing each step to the audit trail, so the organisation can later show '
       'exactly what happened and why.')

# ── 6 Adjacent segments ──────────────────────────────────────────────────────
h1('6', 'Adjacent segments and applications')
para('Each subsection names a real obligation or driver that creates the three conditions of §4, then '
     'maps the prototype’s pattern onto it. The obligations are factual; the mappings are structural.')

h2('6.1', 'Retail and everyday banking')
para('Everyday current-account servicing — standing orders, Direct Debits, disputes, address changes, '
     'account closures — is high-volume, still partly manual, and governed by BCOBS and the Consumer '
     'Duty. The same guided-workflow-plus-cooling-off-plus-CoP pattern applies almost unchanged.')
h2('6.2', 'Wealth and private banking')
para('Discretionary mandates, suitability and complex authorisations demand strong evidence of who '
     'approved what. The dual-authorisation, mandate and audit-trail patterns map directly onto '
     'higher-value, lower-volume wealth processes.')
h2('6.3', 'Corporate and commercial banking')
para('Larger clients bring multi-signatory mandates, treasury payments and heavier KYB. The prototype’s '
     'strictest-rule mandate engine and KYC/KYB flow scale up to corporate authorisation and onboarding.')
h2('6.4', 'Insurance')
para('Insurance conduct rules (ICOBS) and the Consumer Duty require clear, fair, evidenced customer '
     'journeys for quotes, renewals, mid-term adjustments and claims. The guided-workflow and '
     'complaint-handling patterns transfer to policy servicing and claims triage.')
h2('6.5', 'Mortgages and secured lending')
para('MCOB governs affordability, disclosure and cooling-off in mortgage sales and servicing. The '
     'pre-approved-lending flow — offer, disclosure, cooling-off rights, signed draw-down — is the same '
     'shape a mortgage or secured-lending journey needs.')
h2('6.6', 'Cards and payments operations')
para('Card disputes, chargebacks, Direct Debit indemnity claims and payment recalls are evidence-heavy '
     'and time-bound. The Confirmation-of-Payee, Direct-Debit-Guarantee and audit-trail patterns map '
     'onto payments back-office and dispute handling.')
h2('6.7', 'Collections, arrears and vulnerable customers')
para('Consumer Duty and FCA vulnerability guidance require documented, fair treatment of customers in '
     'difficulty. The prototype’s evidence-and-escalation patterns (contact logs, specialist referral, '
     'relationship-manager escalation) map onto collections and vulnerable-customer journeys.')
h2('6.8', 'Onboarding, KYC and financial-crime operations')
para('Onboarding and periodic review run on MLR 2017 / JMLSG obligations and sanctions/PEP screening. '
     'The prototype’s KYC/KYB flow (GOV.UK One Login Lists 1–3, sanctions, PEP, visa) and its '
     'fraud/APP-reimbursement framing map onto financial-crime operations.')
h2('6.9', 'Branch and contact-centre transformation')
para('Any process still handled at a counter or on a call — because it “needs a form” — is a candidate '
     'to become a guided self-serve workflow, with the same control embedded and the same audit trail '
     'produced, reducing cost and posting while improving the evidence base.')

table(
    ['Segment', 'A real obligation / driver', 'Prototype pattern that maps'],
    [['Retail banking','BCOBS · Consumer Duty','Guided servicing + cooling-off + CoP'],
     ['Wealth & private','Suitability, mandate evidence','Dual authorisation + mandate + audit trail'],
     ['Corporate & commercial','Multi-signatory KYB','Strictest-rule mandate engine + KYC/KYB'],
     ['Insurance','ICOBS · Consumer Duty','Guided journeys + DISP complaint handling'],
     ['Mortgages','MCOB affordability & cooling-off','Offer + disclosure + cooling-off + signed draw-down'],
     ['Cards & payments','Dispute / indemnity time limits','CoP + DD Guarantee + evidence trail'],
     ['Collections / vulnerability','Consumer Duty vulnerability guidance','Contact log, escalation, documented fairness'],
     ['Onboarding / financial crime','MLR 2017 / JMLSG · sanctions','KYC/KYB flow + fraud/APP framing'],
     ['Branch / contact centre','Cost, access, Consumer Duty','Self-serve workflow + embedded control + audit']],
    widths=[3.6,4.6,5.1],
    capt='Table 2. Segment-by-segment mapping. Each row names a factual obligation or driver and the '
         'prototype pattern that structurally addresses it.')

# ── 7 Limitations ────────────────────────────────────────────────────────────
h1('7', 'Basis reminder and limitations')
para('The mappings in §6 are claims of applicability, not of readiness. The prototype, as built, is a '
     'front-end demonstration with fictional data and no back-end (see the accompanying Technical '
     'Report). Applying the pattern to any segment above would require the production path of Table 1 — '
     'real services, server-side enforcement, testing and independent assurance — and each segment’s '
     'own regulatory detail. What transfers is the pattern and the evidence discipline, not a finished '
     'system.')

# ── 8 Conclusion ─────────────────────────────────────────────────────────────
h1('8', 'Conclusion')
para('The prototype’s contribution is a repeatable way to make a regulated process paperless without '
     'losing the evidence a regulator expects: guide the task, embed the control, log the outcome. That '
     'pattern is common to retail, wealth, corporate, insurance, mortgages, payments, collections and '
     'financial-crime operations. Business banking was the first place to prove it; it is not the last '
     'place it applies.')

# ── References ───────────────────────────────────────────────────────────────
h1('', 'References')
refs = [
 'Financial Conduct Authority. Consumer Duty (PRIN 2A), 2023.',
 'Financial Conduct Authority. BCOBS — Banking: Conduct of Business Sourcebook.',
 'Financial Conduct Authority. ICOBS — Insurance: Conduct of Business Sourcebook.',
 'Financial Conduct Authority. MCOB — Mortgages and Home Finance: Conduct of Business Sourcebook.',
 'Financial Conduct Authority. DISP — Dispute Resolution: Complaints.',
 'European Banking Authority / FCA. PSD2 Regulatory Technical Standards on SCA, Art. 97.',
 'Payment Systems Regulator. Confirmation of Payee; SI 2019/1215; PS23/3 (APP fraud reimbursement).',
 'Money Laundering Regulations 2017; JMLSG Guidance.',
 'Financial Conduct Authority. SYSC 9 — record-keeping (7-year retention).',
 'D. Berman. Individual Accountability Under the Senior Managers Regime. Macfarlanes / Thomson Reuters, 2016.',
 'Santander Business Banking Prototype — Technical Report, June 2026 (companion document).',
 '1701-uniform/REFERENCE.md — internal technical reference library (112 sections, 121 books & regulatory documents).',
]
for i, r in enumerate(refs, 1):
    p = doc.add_paragraph(); p.paragraph_format.space_after = Pt(2); p.paragraph_format.left_indent = Cm(0.8)
    p.paragraph_format.first_line_indent = Cm(-0.8)
    run = p.add_run(f'[{i}]  '); run.font.bold=True; run.font.size=Pt(8.5); run.font.name='Calibri'; run.font.color.rgb=DARK
    run2 = p.add_run(r); run2.font.size=Pt(8.5); run2.font.name='Calibri'; run2.font.color.rgb=DARK

para('', sa=4)
banner('INTERNAL · CONFIDENTIAL')
para('Santander UK plc · Self-initiated · Completed out of hours, own time · Alan Davidson, Business '
     'Banking Advisor · June 2026', size=8, italic=True, color=LGREY, align=WD_ALIGN_PARAGRAPH.CENTER, sb=6)

out = '/home/user/santander/Santander_Position_Paper.docx'
doc.save(out)
print('Saved →', out)
