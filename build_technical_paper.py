#!/usr/bin/env python3
"""
Santander Business Banking Prototype — Technical Report
Mirrors the layout/reference conventions of the Quantum technical paper:
classification banner · title block · abstract box · keywords · numbered
hierarchical sections · captioned tables · bracketed [n] citations · references
· candid limitations. Branded Santander red, INTERNAL · CONFIDENTIAL.

Produces: Santander_Technical_Report.docx
"""
from docx import Document
from docx.shared import Pt, RGBColor, Inches, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

RED   = RGBColor(0xC8, 0x10, 0x2E)
DARK  = RGBColor(0x1A, 0x1A, 0x1A)
GREY  = RGBColor(0x66, 0x66, 0x66)
LGREY = RGBColor(0x99, 0x99, 0x99)
WHITE = RGBColor(0xFF, 0xFF, 0xFF)

doc = Document()
for s in doc.sections:
    s.page_width, s.page_height = Inches(8.27), Inches(11.69)  # A4
    s.left_margin = s.right_margin = Cm(2.4)
    s.top_margin = s.bottom_margin = Cm(2.2)

# Base style
base = doc.styles['Normal']
base.font.name = 'Calibri'; base.font.size = Pt(10.5); base.font.color.rgb = DARK
base.paragraph_format.space_after = Pt(6); base.paragraph_format.line_spacing = 1.12

def _shade(el, hexv):
    sh = OxmlElement('w:shd'); sh.set(qn('w:val'), 'clear'); sh.set(qn('w:fill'), hexv)
    el.append(sh)

def banner(text):
    t = doc.add_table(rows=1, cols=1); t.alignment = WD_TABLE_ALIGNMENT.CENTER
    c = t.rows[0].cells[0]; _shade(c._tc.get_or_add_tcPr(), '1A1A1A')
    p = c.paragraphs[0]; p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p.add_run(text); r.font.bold = True; r.font.size = Pt(9); r.font.color.rgb = WHITE
    r.font.name = 'Calibri'
    p.paragraph_format.space_after = Pt(0); p.paragraph_format.space_before = Pt(0)

def para(text, size=10.5, bold=False, italic=False, color=DARK, align=WD_ALIGN_PARAGRAPH.LEFT,
         sb=0, sa=6, font='Calibri'):
    p = doc.add_paragraph(); p.alignment = align
    p.paragraph_format.space_before = Pt(sb); p.paragraph_format.space_after = Pt(sa)
    r = p.add_run(text); r.font.size = Pt(size); r.font.bold = bold; r.font.italic = italic
    r.font.color.rgb = color; r.font.name = font
    return p

def h1(num, text):
    p = doc.add_paragraph(); p.paragraph_format.space_before = Pt(14); p.paragraph_format.space_after = Pt(4)
    r = p.add_run(f'{num} {text}'); r.font.size = Pt(14); r.font.bold = True; r.font.color.rgb = DARK; r.font.name='Calibri'
    # red underline rule
    pPr = p._p.get_or_add_pPr(); bdr = OxmlElement('w:pBdr'); b = OxmlElement('w:bottom')
    b.set(qn('w:val'),'single'); b.set(qn('w:sz'),'6'); b.set(qn('w:space'),'4'); b.set(qn('w:color'),'C8102E')
    bdr.append(b); pPr.append(bdr)

def h2(num, text):
    p = doc.add_paragraph(); p.paragraph_format.space_before = Pt(10); p.paragraph_format.space_after = Pt(3)
    r = p.add_run(f'{num} {text}'); r.font.size = Pt(11.5); r.font.bold = True; r.font.color.rgb = RED; r.font.name='Calibri'

def bullet(text):
    p = doc.add_paragraph(style='List Bullet'); p.paragraph_format.space_after = Pt(3)
    p.paragraph_format.left_indent = Cm(0.7)
    r = p.add_run(text); r.font.size = Pt(10.5); r.font.color.rgb = DARK; r.font.name='Calibri'

def abstract_box(text):
    t = doc.add_table(rows=1, cols=1)
    c = t.rows[0].cells[0]; _shade(c._tc.get_or_add_tcPr(), 'F7F2F0')
    # left red border
    tcPr = c._tc.get_or_add_tcPr(); borders = OxmlElement('w:tcBorders')
    lft = OxmlElement('w:left'); lft.set(qn('w:val'),'single'); lft.set(qn('w:sz'),'18'); lft.set(qn('w:space'),'0'); lft.set(qn('w:color'),'C8102E')
    borders.append(lft); tcPr.append(borders)
    p = c.paragraphs[0]
    r = p.add_run('Abstract.  '); r.font.bold = True; r.font.size = Pt(10); r.font.name='Calibri'; r.font.color.rgb=DARK
    r2 = p.add_run(text); r2.font.size = Pt(10); r2.font.name='Calibri'; r2.font.color.rgb=DARK
    p.paragraph_format.space_after = Pt(2)

def caption(text):
    p = doc.add_paragraph(); p.paragraph_format.space_before = Pt(3); p.paragraph_format.space_after = Pt(10)
    r = p.add_run(text); r.font.size = Pt(8.5); r.font.italic = True; r.font.color.rgb = GREY; r.font.name='Calibri'

def table(headers, rows, widths=None, capt=None):
    t = doc.add_table(rows=1, cols=len(headers)); t.alignment = WD_TABLE_ALIGNMENT.CENTER
    t.style = 'Table Grid'
    hdr = t.rows[0].cells
    for i, htext in enumerate(headers):
        _shade(hdr[i]._tc.get_or_add_tcPr(), 'C8102E')
        p = hdr[i].paragraphs[0]; r = p.add_run(htext); r.font.bold=True; r.font.size=Pt(8.5); r.font.color.rgb=WHITE; r.font.name='Calibri'
        p.paragraph_format.space_after = Pt(1)
    for ri, row in enumerate(rows):
        cells = t.add_row().cells
        for ci, val in enumerate(row):
            if ri % 2 == 0: _shade(cells[ci]._tc.get_or_add_tcPr(), 'F7F7F6')
            p = cells[ci].paragraphs[0]; r = p.add_run(val); r.font.size=Pt(8.5); r.font.color.rgb=DARK; r.font.name='Calibri'
            p.paragraph_format.space_after = Pt(1)
    if widths:
        for ci, w in enumerate(widths):
            for row in t.rows:
                row.cells[ci].width = Cm(w)
    if capt: caption(capt)
    return t

# ── Classification banner ────────────────────────────────────────────────────
banner('INTERNAL · CONFIDENTIAL')

# ── Title block ───────────────────────────────────────────────────────────────
para('Santander Business Banking Prototype: A Front-End-First Architecture for '
     'Demonstrating Paperless Workflows with Embedded Regulatory Compliance',
     size=18, bold=True, color=DARK, sb=10, sa=4)
para('Alan Davidson', size=11, bold=True, sa=0)
para('Santander Business Banking', size=10, color=GREY, sa=0)
para('Technical Report · June 2026', size=10, italic=True, color=RED, sa=10)

# Abstract
abstract_box(
    'We present a front-end-only prototype of a Santander Business Banking application that '
    'demonstrates paperless equivalents of branch- and post-based processes — mandate changes, '
    'bulk payments, KYC/KYB, account closures, Making Tax Digital submission, lending, FX, '
    'standing orders and Direct Debits, account unlinking, ring-fencing and complaint handling — '
    'with the relevant regulatory control embedded in each flow rather than bolted on. The system '
    'is a single React component (~6,500 lines) compiled to one self-contained HTML file, runs in '
    'any browser with no backend or real data, and ships an immutable audit trail, dual-authorisation '
    'with personal-PIN signing, working-day-aware cooling-off, and Confirmation of Payee. We describe '
    'the architecture and its deliberate single-file constraint, the entity and mandate models that fork '
    'every compliance path, a feature-to-regulation mapping, and the evaluation the prototype supports. '
    'We state limitations candidly: the data is fictional, there is no production back-end, and the '
    'single-file design is a prototype-stage choice, not a production pattern.')

para('Keywords: paperless banking; dual authorisation; SCA; Confirmation of Payee; FCA Consumer Duty; '
     'Making Tax Digital; cooling-off; Direct Debit Guarantee; FCA DISP complaints; single-file React; '
     'WCAG 2.1 AA; audit trail.', size=9, italic=True, color=GREY, sa=8)

# ── 1 Introduction ─────────────────────────────────────────────────────────────
h1('1', 'Introduction')
para('Business-banking change requests — adding a signatory, paying wages, closing an account, '
     'submitting VAT, cancelling a Direct Debit — have historically meant paper forms, branch visits '
     'and posting documents to a processing centre. Each is slow, error-prone and hard to audit. This '
     'prototype demonstrates that every one of these can be a guided, self-serve digital workflow in '
     'which the regulatory control is part of the flow, not a separate compliance step [1,2].')
para('The prototype is deliberately front-end only. It holds no real customer data, makes no API calls '
     '(other than Google Fonts), and is built to be opened on any device for internal presentations and '
     'customer research. This scopes the contribution precisely: it is a demonstration of workflow design '
     'and embedded compliance, not a production banking system.')

h2('1.1', 'Contributions')
bullet('A single-file React architecture (~6,500 lines) that compiles to one self-contained ~890 KB '
       'HTML file (~194 KB gzip), readable top-to-bottom and deployable as a static asset.')
bullet('Thirteen step-based workflow wizards spanning the key business-banking operations, each forking '
       'by entity type and signing rule.')
bullet('Embedded controls: dual authorisation with personal-PIN signing (SCA), working-day-aware 24-hour '
       'cooling-off, Confirmation of Payee, KYC/KYB, and an FCA DISP-compliant complaint flow.')
bullet('A seven-entity model and three-tier mandate engine (Any-1 / Any-2 / All) that drives whether a '
       'workflow needs co-signers or a cooling-off period.')
bullet('An immutable, actor-and-timestamp audit trail aligned to FCA SYSC 9 — the demonstrability '
       'evidence the Senior Managers Regime expects [3].')
bullet('A consistent design system (WCAG 2.1 AA contrast, focus-visible, tabular figures) enforced by a '
       'six-station pre-commit review process.')

# ── 2 Background ───────────────────────────────────────────────────────────────
h1('2', 'Background and Regulatory Context')
h2('2.1', 'Strong Customer Authentication and dual authorisation')
para('Payment and mandate actions are authorised under PSD2 Regulatory Technical Standards Art. 97 [4]. '
     'The prototype models this as a per-account signing rule and a personal-PIN signing step, with a '
     'six-tier step-up matrix escalating to biometric and Voice ID for higher-risk actions.')
h2('2.2', 'Confirmation of Payee and payment safety')
para('Before any payment leaves, a Confirmation of Payee check verifies the payee name against the '
     'account, with an explicit override path that cannot be skipped silently (PSR / SI 2019/1215) [5]. '
     'APP-fraud reimbursement obligations under PSR PS23/3 frame the complaint and fraud paths [6].')
h2('2.3', 'Cooling-off, Consumer Duty and complaints')
para('High-impact actions carry a 24-working-hour cooling-off period (FCA BCOBS 4A) during which the '
     'customer can cancel without penalty; the timer pauses over weekends and bank holidays. The '
     'complaint workflow follows FCA DISP — eligible-complainant checks, escalation triage and FOS '
     'signposting — consistent with the Consumer Duty obligation to deliver good outcomes [1,7].')
h2('2.4', 'The single-file architecture decision')
para('A central design question separates a defensible prototype from a fragile one. The obvious answer '
     'to "make it production-shaped" — split into modules behind a backend — is the wrong one for a '
     'demonstration whose value is being read and understood end-to-end. We therefore keep the whole '
     'product in one file, accept the trade-off explicitly (§7), and document the production path (§7) '
     'rather than pretend the prototype is the product.')

# ── 3 System overview ──────────────────────────────────────────────────────────
h1('3', 'System Overview and Design Principles')
para('The prototype is organised around three principles. Front-end by default: it runs with no network '
     'access and no backend, so it can be demonstrated anywhere. One readable artefact: the entire product '
     'is a single React component, so a reviewer can follow every workflow without navigating modules. '
     'Compliance in the flow: each control is rendered as part of the workflow that needs it, with the '
     'governing regulation named in the UI.')
caption('Figure 1. The prototype is one React component compiled to a single HTML file; a shared state '
        'tree feeds two render shells (mobile bottom-nav and desktop sidebar) that present the same '
        'screens and workflow overlays.')

# ── 4 Architecture ─────────────────────────────────────────────────────────────
h1('4', 'Application Architecture')
h2('4.1', 'Single-component structure')
para('App.jsx is one function component. All hooks are declared at the top (120+ state variables) before '
     'any logic, because the sub-components — workflow renderers, screens and sheets — are closures and '
     'cannot hold their own hooks. A reset function (closeWorkflow) clears all workflow state to prevent '
     'ghost state between flows.')
h2('4.2', 'Navigation model')
para('Three state layers compose every view: a tab (home, approve, audit, MTD, statements), a workflow '
     'overlay (one of thirteen, or none), and a step index within the active workflow. A viewMode flag '
     'switches the mobile and desktop shells while preserving all in-progress state.')
h2('4.3', 'Entity system')
para('A static map configures seven entity types; the active type re-derives labels, required documents, '
     'registration numbers, the accounts list and mandate defaults (Table 1).')
table(
    ['Entity type', 'Mandate default', 'Principal', 'Companies House', 'Board minutes'],
    [['Sole trader','Any-1','Sole trader','No','No'],
     ['Partnership','Any-1 / Any-2','Partner','No','On rename (RM)'],
     ['Limited company','Any-2','Director','Yes','No'],
     ['LLP','Any-2','Member','Yes (LLP no.)','No'],
     ['Registered charity','All','Trustee','Charity Comm.','Yes — all changes'],
     ['Club','All','Committee member','No','Yes — all changes'],
     ['Society','All','Committee member','No','Yes — all changes']],
    widths=[3.6,2.6,2.6,2.8,3.0],
    capt='Table 1. The seven entity types and the configuration each forks. The mandate default drives '
         'whether a workflow requires co-signers and a cooling-off period.')
h2('4.4', 'Mandate rules')
para('Each account carries a signing rule — Any-1, Any-2 or All. A helper, getMandateFor(accounts), '
     'selects the strictest rule across the selected accounts (All > Any-2 > Any-1). The result determines '
     'whether a workflow requires a co-signer and/or triggers the 24-hour cooling-off period.')
h2('4.5', 'Design system')
para('Tailwind utilities with a brand palette (#c8102e red, #faf6ef warm background, stone neutrals) and '
     'an inline CSS layer for brand cards and animations. The system enforces WCAG 2.1 AA contrast, a '
     'visible focus ring on every interactive element, one primary call-to-action per view, and tabular '
     'figures on all monetary amounts [8].')

# ── 5 Workflows ────────────────────────────────────────────────────────────────
h1('5', 'Workflows')
para('Thirteen step-based wizards render on top of the main screens; each reads and writes the shared '
     'state tree and forks by entity and signing rule. Table 2 catalogues them.')
table(
    ['Workflow', 'Embedded control', 'Regulatory basis'],
    [['Account closure','Destination CoP, credit check, cooling-off','FCA BCOBS 4A'],
     ['Partner-unreachable escalation','4 branches, evidence, restriction','FCA PS22/9'],
     ['Mandate changes','KYC/KYB, board minutes, signing rule','MLR 2017 / JMLSG'],
     ['Bulk payments / wages','CoP per payee, dual authorisation','PSR; PSD2 RTS'],
     ['Business details update','Companies House sync','Companies House'],
     ['Dormant reactivation','FSCS notice, 12-month threshold','FSCS'],
     ['MTD VAT submission','Auto-categorise, VAT100, HMRC API','HMRC MTD'],
     ['Personal/business unlink','Render-tree removal, CRM update','GDPR Art.5(1)(c)'],
     ['Credit ring-fence','Persisted formal instruction','GDPR Art.5(1)(c)'],
     ['Pre-approved lending','14-day cooling-off, repayment calc','CCA 1974'],
     ['International FX payment','Fee disclosure, MLR screening','MLR 2017'],
     ['Standing orders & Direct Debits','DD Guarantee, CoP on set-up','PSR; SI 2019/1215'],
     ['Complaint handling','Eligibility, triage, FOS signpost','FCA DISP']],
    widths=[4.4,5.3,3.6],
    capt='Table 2. The thirteen workflow wizards, the control embedded in each, and its regulatory basis.')
h2('5.1', 'Dual authorisation and signing')
para('Pending actions reach a signature queue. Signing uses a personal four-digit PIN by default '
     '(biometric and Voice ID remain selectable), framed as SCA. A signed payment then enters a '
     'ten-second cancel window before release.')
h2('5.2', 'Cooling-off')
para('The cooling-off timer adds 24 working hours, advancing hour-by-hour and only decrementing on '
     'weekday, non-bank-holiday hours; a request started on a Friday therefore executes the following '
     'Monday. A live card shows a progress bar, with Cancel and Do-it-now controls.')

# ── 6 Compliance mapping ─────────────────────────────────────────────────────────
h1('6', 'Compliance Mapping')
para('Each control is tied to a specific obligation rather than a vague claim (Table 3).')
table(
    ['Control', 'Regulation', 'Implementation'],
    [['Cooling-off','FCA BCOBS 4A','Working-hours aware, cancellable, live progress'],
     ['Dual authorisation','PSD2 RTS Art.97','Any-1/2/All per account, PIN signing, co-sign window'],
     ['KYC / KYB','MLR 2017 / JMLSG','GOV.UK One Login Lists 1–3, sanctions, PEP, visa'],
     ['Audit trail','FCA SYSC 9','7-year, append-only, actor + timestamp per event'],
     ['Confirmation of Payee','PSR / SI 2019/1215','Pre-payment name check, explicit override'],
     ['Complaints','FCA DISP','D+3/D+56 deadlines, eligibility, FOS signpost'],
     ['Data separation','GDPR Art.5(1)(c)','Render-tree removal, CRM unlink, credit ring-fence'],
     ['Accessibility','WCAG 2.1 AA','Contrast ≥4.5:1, focus-visible, tabular figures']],
    widths=[3.4,3.6,6.3],
    capt='Table 3. Feature-to-regulation mapping. No control is included without a named obligation behind it.')

# ── 7 Limitations & production path ──────────────────────────────────────────────
h1('7', 'Limitations and Path to Production')
para('We state the boundaries plainly, in the spirit of a defensible rather than promotional account.')
bullet('Fictional data only. All names, account numbers, balances and amounts are invented; there is no '
       'real customer data and no money movement.')
bullet('No back-end. The prototype is front-end only; the workflows imply but do not call production '
       'services (payments, KYC, HMRC, audit log). A Node.js service layer is the natural Phase-2 step [9].')
bullet('Single-file by design. The ~6,500-line single component is intentional for a readable prototype, '
       'not a production structure; hardening would migrate to TypeScript and a modular architecture [10].')
bullet('Demonstration of controls, not certification. Embedded controls model the correct regulatory '
       'behaviour for research and presentation; they are not a substitute for a controlled, tested and '
       'independently assured production implementation.')

# ── 8 Conclusion ─────────────────────────────────────────────────────────────────
h1('8', 'Conclusion')
para('The prototype shows that the paper-and-branch processes of business banking can be expressed as '
     'guided digital workflows with the governing regulatory control built into each flow, in a single '
     'readable, portable artefact. It is a faithful demonstration of design and embedded compliance — and '
     'an explicit, documented foundation for a production system, not a stand-in for one.')

# ── References ────────────────────────────────────────────────────────────────────
h1('', 'References')
refs = [
 'Financial Conduct Authority. Consumer Duty (PRIN 2A), 2023.',
 'Financial Conduct Authority. BCOBS — Banking: Conduct of Business Sourcebook.',
 'D. Berman. Individual Accountability Under the Senior Managers Regime. Macfarlanes / Thomson Reuters, 2016.',
 'European Banking Authority / FCA. PSD2 Regulatory Technical Standards on Strong Customer Authentication, Art. 97.',
 'Payment Systems Regulator. Confirmation of Payee; SI 2019/1215.',
 'Payment Systems Regulator. PS23/3 — APP fraud mandatory reimbursement, 2024.',
 'Financial Conduct Authority. DISP — Dispute Resolution: Complaints.',
 'W3C. Web Content Accessibility Guidelines (WCAG) 2.1.',
 'R. Kumar. Ultimate Node.js for Cross-Platform App Development; N. Murray. Fullstack Node.js.',
 'M. Pocock & T. Bell. Total TypeScript: The Essentials.',
 'HMRC. Making Tax Digital for VAT — API v1.0.',
 'Money Laundering Regulations 2017; JMLSG Guidance.',
 'Financial Conduct Authority. SYSC 9 — record-keeping (7-year retention).',
 '1701-uniform/REFERENCE.md — internal technical reference library (112 sections, 121 books & regulatory documents).',
]
for i, r in enumerate(refs, 1):
    p = doc.add_paragraph(); p.paragraph_format.space_after = Pt(2); p.paragraph_format.left_indent = Cm(0.8)
    p.paragraph_format.first_line_indent = Cm(-0.8)
    run = p.add_run(f'[{i}]  '); run.font.bold = True; run.font.size = Pt(8.5); run.font.name='Calibri'; run.font.color.rgb=DARK
    run2 = p.add_run(r); run2.font.size = Pt(8.5); run2.font.name='Calibri'; run2.font.color.rgb=DARK

# Footer note
para('', sa=4)
banner('INTERNAL · CONFIDENTIAL')
para('Santander UK plc · Self-initiated · Completed out of hours, own time · Alan Davidson, Business '
     'Banking Advisor · June 2026', size=8, italic=True, color=LGREY, align=WD_ALIGN_PARAGRAPH.CENTER, sb=6)

out = '/home/user/santander/Santander_Technical_Report.docx'
doc.save(out)
print('Saved →', out)
