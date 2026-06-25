"""
Santander Complaint Writing Guide
Produces: Santander_Complaint_Guide.docx
"""
from docx import Document
from docx.shared import Pt, RGBColor, Inches, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.style import WD_STYLE_TYPE
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
import copy

doc = Document()

# ── Page setup ────────────────────────────────────────────────────────────────
section = doc.sections[0]
section.page_width  = Inches(8.27)   # A4
section.page_height = Inches(11.69)
section.left_margin   = Inches(1.0)
section.right_margin  = Inches(1.0)
section.top_margin    = Inches(0.9)
section.bottom_margin = Inches(0.9)

RED   = RGBColor(0xC8, 0x10, 0x2E)
DARK  = RGBColor(0x1C, 0x19, 0x17)
GREY  = RGBColor(0x78, 0x71, 0x6C)
WHITE = RGBColor(0xFF, 0xFF, 0xFF)

# ── Style helpers ─────────────────────────────────────────────────────────────
def set_font(run, size=11, bold=False, color=None, italic=False):
    run.font.name = 'Calibri'
    run.font.size = Pt(size)
    run.font.bold = bold
    run.font.italic = italic
    if color:
        run.font.color.rgb = color

def shade_paragraph(para, hex_fill):
    pPr = para._p.get_or_add_pPr()
    shd = OxmlElement('w:shd')
    shd.set(qn('w:val'), 'clear')
    shd.set(qn('w:color'), 'auto')
    shd.set(qn('w:fill'), hex_fill)
    pPr.append(shd)

def h1(text):
    p = doc.add_paragraph()
    shade_paragraph(p, 'C8102E')
    p.paragraph_format.space_before = Pt(18)
    p.paragraph_format.space_after  = Pt(4)
    p.paragraph_format.left_indent  = Cm(0.3)
    run = p.add_run(text)
    set_font(run, size=16, bold=True, color=WHITE)
    return p

def h2(text):
    p = doc.add_paragraph()
    shade_paragraph(p, 'F5F5F4')
    p.paragraph_format.space_before = Pt(14)
    p.paragraph_format.space_after  = Pt(2)
    p.paragraph_format.left_indent  = Cm(0.2)
    run = p.add_run(text)
    set_font(run, size=13, bold=True, color=DARK)
    return p

def h3(text):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(10)
    p.paragraph_format.space_after  = Pt(1)
    run = p.add_run(text)
    set_font(run, size=11, bold=True, color=RED)
    return p

def label(text):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(7)
    p.paragraph_format.space_after  = Pt(1)
    run = p.add_run(text)
    set_font(run, size=10, bold=True, color=DARK)
    return p

def body(text, italic=False):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(0)
    p.paragraph_format.space_after  = Pt(4)
    p.paragraph_format.left_indent  = Cm(0.5)
    run = p.add_run(text)
    set_font(run, size=10, italic=italic, color=DARK)
    return p

def note(text):
    p = doc.add_paragraph()
    shade_paragraph(p, 'FEF9F0')
    p.paragraph_format.space_before = Pt(4)
    p.paragraph_format.space_after  = Pt(4)
    p.paragraph_format.left_indent  = Cm(0.3)
    run = p.add_run('Note: ' + text)
    set_font(run, size=9, italic=True, color=GREY)
    return p

def letter_block(text):
    p = doc.add_paragraph()
    shade_paragraph(p, 'EFF6FF')
    p.paragraph_format.space_before = Pt(3)
    p.paragraph_format.space_after  = Pt(3)
    p.paragraph_format.left_indent  = Cm(0.5)
    p.paragraph_format.right_indent = Cm(0.5)
    run = p.add_run(text)
    set_font(run, size=10, italic=True, color=DARK)
    return p

def divider():
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(2)
    p.paragraph_format.space_after  = Pt(2)
    run = p.add_run('─' * 80)
    set_font(run, size=7, color=GREY)
    return p

def section_label(number, title):
    p = doc.add_paragraph()
    shade_paragraph(p, 'F5F5F4')
    p.paragraph_format.space_before = Pt(8)
    p.paragraph_format.space_after  = Pt(2)
    p.paragraph_format.left_indent  = Cm(0.3)
    r1 = p.add_run(f'{number} — ')
    set_font(r1, size=10, bold=True, color=RED)
    r2 = p.add_run(title)
    set_font(r2, size=10, bold=True, color=DARK)
    return p

# ── Complaint type helper ─────────────────────────────────────────────────────
def write_stage1(s1, s2, s3, s4, s5_text, letter_type, letter_para1, letter_para2, letter_para3=None):
    section_label('1', 'What is the customer saying has happened?')
    body(s1)
    section_label('2', 'What have you done to investigate the issue raised by the customer?')
    body(s2)
    section_label('3', 'What actions have you taken for this customer as a result of their complaint?')
    body(s3)
    section_label('4', 'Why did you take the action detailed above?')
    body(s4)
    section_label('5', 'Payments:')
    body(s5_text)
    h3(f'Letter — {letter_type}')
    letter_block(letter_para1)
    letter_block(letter_para2)
    if letter_para3:
        letter_block(letter_para3)

def write_escalation(e1, e2, e3, e4_text, letter_type, letter_para):
    section_label('1', 'Why has the customer come back to discuss this complaint again?')
    body(e1)
    section_label('2', 'What actions have you completed for the customer?')
    body(e2)
    section_label('3', 'What changes have you made to the resolution, and why?')
    body(e3)
    section_label('4', 'Payments:')
    body(e4_text)
    h3(f'Letter — {letter_type}')
    letter_block(letter_para)

# ══════════════════════════════════════════════════════════════════════════════
# COVER PAGE
# ══════════════════════════════════════════════════════════════════════════════
p = doc.add_paragraph()
shade_paragraph(p, 'C8102E')
p.paragraph_format.space_before = Pt(40)
p.paragraph_format.space_after  = Pt(0)
p.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = p.add_run('Santander Business Banking')
set_font(run, size=22, bold=True, color=WHITE)

p = doc.add_paragraph()
shade_paragraph(p, 'C8102E')
p.paragraph_format.space_before = Pt(0)
p.paragraph_format.space_after  = Pt(40)
p.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = p.add_run('Complaint Writing Guide')
set_font(run, size=16, bold=False, color=WHITE)

p = doc.add_paragraph()
p.paragraph_format.space_before = Pt(14)
p.paragraph_format.space_after  = Pt(4)
p.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = p.add_run('Stage 1 and Escalation — Templates and Worked Examples')
set_font(run, size=12, bold=True, color=DARK)

p = doc.add_paragraph()
p.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = p.add_run('Internal use only · Complaint handling reference · June 2026')
set_font(run, size=10, italic=True, color=GREY)

doc.add_page_break()

# ══════════════════════════════════════════════════════════════════════════════
# SECTION A — FRAMEWORK
# ══════════════════════════════════════════════════════════════════════════════
h1('A — Complaint Form Structure')

h2('Stage 1 — Five Sections')
note('Sections 1 and 3 must be written so the text can be pulled directly into a letter template.')
section_label('1', 'What is the customer saying has happened?')
body('The customer\'s claim. Write in plain, third-person language: what happened, when, what was lost, what they are seeking. No internal jargon.')
section_label('2', 'What have you done to investigate the issue raised by the customer?')
body('Your evidence. What records, systems, or third-party confirmations you reviewed and what they showed. Be specific about dates and findings.')
section_label('3', 'What actions have you taken for this customer as a result of their complaint?')
body('The resolution. What you did to fix the issue and any compensation offered. Write so it can be lifted straight into the letter.')
section_label('4', 'Why did you take the action detailed above?')
body('Your rationale. Santander error or no Santander error. What caused the error or why no error was found. Include goodwill reasoning only if goodwill was given.')
section_label('5', 'Payments:')
body('Compensation breakdown. Figure, calculation method, goodwill element and rationale, total. If no compensation: state why.')

h2('Escalation — Four Sections')
section_label('1', 'Why has the customer come back to discuss this complaint again?')
body('What aspect of the Stage 1 decision the customer disputed and any new evidence they provided.')
section_label('2', 'What actions have you completed for the customer?')
body('What additional records were reviewed on escalation and what decision was reached.')
section_label('3', 'What changes have you made to the resolution, and why?')
body('If changed: what new evidence justified the change and what the new resolution is. If no change: why the original decision stands.')
section_label('4', 'Payments:')
body('Revised calculation if compensation changed, or confirm original offer stands.')

h1('B — Letter Templates')

h2('Upheld — Stage 1')
letter_block("I'm sorry you've had to contact us about [what the complaint is about].")
letter_block("I've looked into what happened and I can see the issue(s) you've raised was / were caused by [root cause]. I understand this has caused you [impact — inconvenience / worry / upset / stress — use customer's own language], and for this I'm truly sorry.")
letter_block("To fix this for you I've [resolution — what action was taken and how the issue was fixed].")

h2('Declined — Stage 1')
letter_block("I'm sorry you've had to contact us about [what the complaint is about].")
letter_block("I understand this situation has caused you problems, but I've looked into your case and this wasn't because of a mistake we made. This is because [brief one-line reason].")
letter_block("To explain this further [full explanation of the evidence found and why Santander was not at fault].")

h2('Escalation — Resolution Changed')
letter_block("I've looked at the details of your complaint and I agree that our original decision [explain why the previous decision was incorrect and what actions have been taken to resolve it].")

h2('Escalation — No Change in Original Resolution')
letter_block("I'm sorry you don't agree with the resolution previously offered about [brief explanation of what the complaint was about].")
letter_block("Taking all the circumstances of your complaint into consideration, I believe our original offer of £[AMOUNT] is fair and reasonable.\n\nOR — reiterate the original decision, e.g.: The delay in [X] happened after [explanation].")

h1('C — Interest Compensation Formula')
body('Compensation = Principal × (Annual Rate ÷ 100) ÷ 365 × Number of calendar days delayed')
body('Example: £8,500 at 3.5% for 10 days = £8,500 × 0.035 ÷ 365 × 10 = £8.15')
note('Use the rate applicable during the delay period. Round to the nearest penny. Document in section 5.')

h1('D — FCA DISP Quick Reference')
table = doc.add_table(rows=1, cols=2)
table.style = 'Table Grid'
hdr = table.rows[0].cells
hdr[0].text = 'Rule'
hdr[1].text = 'Requirement'
for cell in hdr:
    for para in cell.paragraphs:
        for run in para.runs:
            run.font.bold = True
            run.font.size = Pt(9)
            run.font.name = 'Calibri'

rows = [
    ('DISP 1.3', 'Complaint must be acknowledged promptly'),
    ('DISP 1.6.1R — D+3', 'Summary resolution within 3 business days (no letter required if customer accepts)'),
    ('DISP 1.6.2R — D+56', 'Full investigation: final response within 8 weeks (56 days)'),
    ('DISP 1.6.4R', 'If no final response by D+56: write to customer explaining delay and right to go to FOS'),
    ('DISP 2.7 — Eligible complainants', 'Individuals, micro-enterprises (< 10 staff / < £2m turnover), charities < £1m income, small trusts'),
    ('DISP 2.8.2R — Time bar', '>6 years from event, or >3 years from when customer reasonably should have known'),
    ('DISP 1.10', 'Half-yearly aggregate complaints data return to FCA'),
    ('FOS referral', 'Customer has 6 months from final response to refer — must be stated in every final response letter'),
    ('FOS contact', '0800 023 4567 · financial-ombudsman.org.uk'),
]
for rule, req in rows:
    row = table.add_row().cells
    row[0].text = rule
    row[1].text = req
    for cell in row:
        for para in cell.paragraphs:
            for run in para.runs:
                run.font.size = Pt(9)
                run.font.name = 'Calibri'

doc.add_page_break()

# ══════════════════════════════════════════════════════════════════════════════
# COMPLAINT TYPE 1 — ISA TRANSFER DELAY
# ══════════════════════════════════════════════════════════════════════════════
h1('1 — ISA Transfer Delay')
note('Cash ISA industry standard: 15 working days. Applies whether transfer is to or from Santander.')

h2('1A — Stage 1 · Upheld (Santander processing error)')
write_stage1(
    s1="The customer states that they requested the transfer of their Cash ISA to Santander on [DATE OF REQUEST]. They advise that the transfer was not completed within the expected 15 working day timeframe and that as of [DATE CUSTOMER CONTACTED US] their funds of £[AMOUNT] had not been credited to their Santander Cash ISA. The customer states that as a result of this delay they have lost interest on their savings during the period the funds were unavailable and are seeking resolution and compensation.",
    s2="I have reviewed the customer's ISA account and confirmed that a transfer request was submitted on [DATE]. The expected completion date under the 15 working day industry standard was [DATE]. I reviewed our internal ISA operations logs and confirmed that a processing error on our side caused the transfer instruction to fail and re-enter the queue without the customer being notified. I confirmed the transfer completed on [ACTUAL COMPLETION DATE], representing a delay of [X] working days beyond the expected timeframe. I calculated the interest lost on £[AMOUNT] over the delay period at the customer's agreed ISA rate of [RATE]%.",
    s3="I have confirmed to the customer that their ISA transfer is now complete and their funds of £[AMOUNT] are held in their Santander Cash ISA as of [DATE]. I have calculated the interest lost during the delay period and arranged a compensation payment to cover this, along with a goodwill gesture in recognition of the inconvenience and distress caused.",
    s4="This complaint has been upheld as a Santander error. Our investigation confirmed that an internal processing error caused the transfer instruction to fail and re-enter our processing queue without notification to the customer. This resulted in the transfer completing [X] working days beyond the 15 working day industry standard. As the delay was caused by our error I have compensated the customer for the interest lost during the delay period. I have also offered a goodwill gesture in recognition of the worry and inconvenience this caused.",
    s5_text="Lost interest: £[AMOUNT] × [RATE]% ÷ 365 × [X working days] = £[CALCULATED INTEREST]\nGoodwill gesture: £[AMOUNT] — in recognition of the inconvenience and distress caused by our error\nTotal: £[TOTAL]",
    letter_type="Upheld Stage 1",
    letter_para1="I'm sorry you've had to contact us about the delay in transferring your Cash ISA to Santander.",
    letter_para2="I've looked into what happened and I can see the issue you've raised was caused by an internal processing error on our side, which meant your transfer instruction failed and re-entered our queue without us notifying you. I understand this has caused you worry about the whereabouts of your savings and inconvenience during the delay, and for this I'm truly sorry.",
    letter_para3="To fix this for you I've confirmed your ISA transfer is now complete and your funds of £[AMOUNT] are in your Santander Cash ISA. I've also arranged a payment of £[TOTAL] to cover the interest you lost during the delay and to recognise the inconvenience this has caused you.",
)

divider()
h2('1B — Stage 1 · Declined (previous provider caused the delay)')
write_stage1(
    s1="The customer states that they requested the transfer of their Cash ISA from [PREVIOUS PROVIDER] to Santander on [DATE OF REQUEST]. They advise the transfer was not completed within the expected 15 working day timeframe and are holding Santander responsible for the delay. The customer is requesting compensation for the inconvenience and any interest lost during the delay period.",
    s2="I have reviewed the customer's ISA account and confirmed that a transfer request was submitted on [DATE], with an expected completion date of [DATE]. I reviewed our internal ISA operations records and confirmed that Santander submitted the transfer instruction to [PREVIOUS PROVIDER] within the required timeframe on [DATE]. I obtained confirmation from our ISA operations team that [PREVIOUS PROVIDER] did not respond to or action the transfer request within the 15 working day industry standard. Santander was not responsible for any part of the delay. The transfer completed on [DATE] following our team chasing [PREVIOUS PROVIDER].",
    s3="I have confirmed to the customer that their ISA transfer is now complete. I have explained that the delay was caused by [PREVIOUS PROVIDER] failing to release the funds within the required 15 working day timeframe and that Santander acted correctly and within our obligations throughout. I have advised the customer to raise the matter directly with [PREVIOUS PROVIDER] as the responsible party and signposted them to the Financial Ombudsman Service should they remain dissatisfied.",
    s4="This complaint has not been upheld as there is no Santander error. Our investigation confirmed Santander submitted the transfer instruction to [PREVIOUS PROVIDER] within the required timeframe and that the delay was caused solely by [PREVIOUS PROVIDER] failing to act within the 15 working day industry standard. As Santander did not cause the delay we are not in a position to offer compensation. The customer has been directed to [PREVIOUS PROVIDER] to seek redress.",
    s5_text="No compensation offered — the delay was not caused by a Santander error. Customer directed to [PREVIOUS PROVIDER] to seek redress for the delay period.",
    letter_type="Declined Stage 1",
    letter_para1="I'm sorry you've had to contact us about the delay in transferring your Cash ISA to Santander.",
    letter_para2="I understand this situation has caused you problems, but I've looked into your case and this wasn't because of a mistake we made. This is because the delay was caused by your previous provider, not by Santander.",
    letter_para3="To explain this further, I can confirm that Santander submitted the transfer instruction to [PREVIOUS PROVIDER] on [DATE], within the required timeframe. Our records show that [PREVIOUS PROVIDER] did not release your funds within the 15 working day industry standard, which caused the delay you experienced. As we were not responsible for this delay we are unable to offer compensation on this occasion. I would recommend contacting [PREVIOUS PROVIDER] directly to raise a complaint and seek any compensation for interest lost.",
)

divider()
h2('1C — Escalation · Resolution Changed (new evidence — Santander error identified)')
write_escalation(
    e1="The customer disagrees with the original decision to decline their complaint regarding the delay in their Cash ISA transfer. They maintain Santander bears some responsibility and have provided correspondence from [PREVIOUS PROVIDER] confirming the date the funds were released, which they believe contradicts the basis of the original decision.",
    e2="On escalation I conducted a more detailed review of the ISA transfer records, including transfer instruction timestamps in our internal systems. I reviewed the correspondence provided by the customer from [PREVIOUS PROVIDER] confirming the date the funds were released. This identified that [PREVIOUS PROVIDER] released the funds on [DATE], within the required 15 working day timeframe, and that the subsequent delay of [X] working days occurred within Santander's own processing systems. On the basis of this new evidence I have changed the decision from declined to upheld.",
    e3="I have changed the resolution from declined to upheld. The original decision assumed the delay was entirely caused by [PREVIOUS PROVIDER]. The escalation review identified that [PREVIOUS PROVIDER] released the funds within the required timeframe and the delay occurred within Santander's own processing. As a Santander error has now been identified I have arranged compensation for the interest lost and a goodwill gesture to recognise both the delay and the inconvenience of the incorrect initial decision.",
    e4_text="Lost interest: £[AMOUNT] × [RATE]% ÷ 365 × [X working days] = £[CALCULATED INTEREST]\nGoodwill gesture: £[AMOUNT] — original error plus incorrect initial decision\nTotal: £[TOTAL]",
    letter_type="Escalation · Resolution Changed",
    letter_para="I've looked at the details of your complaint and I agree that our original decision was incorrect. Having reviewed your case in more detail and considered the correspondence you provided from [PREVIOUS PROVIDER], I can see that your previous provider released your funds within the required timeframe and the delay was caused by an error in our own processing. I have therefore changed our decision and upheld your complaint. I have arranged a payment of £[TOTAL] to cover the interest you lost during the delay and to recognise the inconvenience caused by both the original error and our initial incorrect decision.",
)

doc.add_page_break()

# ══════════════════════════════════════════════════════════════════════════════
# COMPLAINT TYPE 2 — INCORRECT CHARGE / FEE
# ══════════════════════════════════════════════════════════════════════════════
h1('2 — Incorrect Charge or Fee Applied')
note('Common causes: duplicate charge, fee applied after account type change, tariff error, charges not matching agreed terms.')

h2('2A — Stage 1 · Upheld (charge incorrectly applied)')
write_stage1(
    s1="The customer states that a charge of £[AMOUNT] was applied to their account on [DATE] and that they were not advised this charge would be applied and/or that it does not reflect the agreed terms on their account. The customer is requesting a full refund of the charge and an explanation of how it was applied.",
    s2="I have reviewed the customer's account and the charge applied on [DATE]. I reviewed the account tariff and terms applicable at the time and confirmed that the charge of £[AMOUNT] was applied in error — [the account type does not attract this charge / the charge was applied twice / the customer's account was moved to a new tariff without proper notification]. I confirmed the charge was not in line with the terms and conditions applicable to this account.",
    s3="I have confirmed to the customer that the charge of £[AMOUNT] applied on [DATE] was applied in error and have arranged a full refund to their account. I have also checked the account for any similar charges applied in the same period and confirmed [no further charges were incorrectly applied / a further £[AMOUNT] has been identified and refunded]. The total refund of £[TOTAL REFUND] will be credited within [X] working days.",
    s4="This complaint has been upheld as a Santander error. Our investigation confirmed the charge was applied incorrectly because [specific reason]. As this was our error I have arranged a full refund of the charge. [I have also offered a goodwill gesture of £[AMOUNT] in recognition of the inconvenience caused / No goodwill has been offered as the error has been fully remedied by the refund.]",
    s5_text="Refund: £[AMOUNT] (incorrect charge)\n[Additional refund: £[AMOUNT] if further charges found]\n[Goodwill gesture: £[AMOUNT] — inconvenience caused]\nTotal: £[TOTAL]",
    letter_type="Upheld Stage 1",
    letter_para1="I'm sorry you've had to contact us about the charge applied to your account.",
    letter_para2="I've looked into what happened and I can see the charge of £[AMOUNT] applied on [DATE] was caused by [root cause — e.g. a tariff error on our side]. I understand this has caused you frustration and concern about your account charges, and for this I'm truly sorry.",
    letter_para3="To fix this for you I've arranged a full refund of £[TOTAL REFUND] back to your account, which you should see within [X] working days.",
)

divider()
h2('2B — Stage 1 · Declined (charge correctly applied per agreed terms)')
write_stage1(
    s1="The customer states that a charge of £[AMOUNT] was applied to their account on [DATE] and that they were not expecting this charge and do not agree it should have been applied. The customer is requesting a full refund.",
    s2="I have reviewed the customer's account and the charge applied on [DATE]. I reviewed the account terms and conditions applicable at the time and confirmed that the charge of £[AMOUNT] was applied correctly in line with the agreed tariff for this account type. I reviewed the customer's account opening documentation and confirmed that the relevant charges were disclosed at the time the account was opened / at the time the tariff changed. I can confirm no error was made in applying this charge.",
    s3="I have explained to the customer that the charge was applied correctly in line with their agreed account terms and that we are therefore unable to offer a refund. I have taken the opportunity to review the customer's account and have provided information on account options that may better suit their needs and reduce future charges.",
    s4="This complaint has not been upheld as there is no Santander error. Our investigation confirmed the charge of £[AMOUNT] was applied correctly in accordance with the terms and conditions applicable to the customer's account. The charge was disclosed at account opening / tariff change. As no error has been made we are not in a position to offer a refund.",
    s5_text="No refund offered — the charge was applied correctly per the agreed account terms.",
    letter_type="Declined Stage 1",
    letter_para1="I'm sorry you've had to contact us about the charge applied to your account.",
    letter_para2="I understand this situation has caused you concern, but I've looked into your case and this wasn't because of a mistake we made. This is because the charge was applied correctly in line with your agreed account terms.",
    letter_para3="To explain this further, I can confirm that the charge of £[AMOUNT] applied on [DATE] is in line with the tariff for your account type, which was provided to you when [your account was opened / your account moved to this tariff] on [DATE]. As the charge has been correctly applied we are unable to offer a refund on this occasion.",
)

doc.add_page_break()

# ══════════════════════════════════════════════════════════════════════════════
# COMPLAINT TYPE 3 — PAYMENT ERROR (WRONG AMOUNT / WRONG ACCOUNT)
# ══════════════════════════════════════════════════════════════════════════════
h1('3 — Payment Error (Wrong Amount or Wrong Account)')
note('Covers: bank-initiated errors (direct debit collected incorrectly, standing order wrong amount) and misdirected payments where Santander\'s system contributed.')

h2('3A — Stage 1 · Upheld (Santander payment processing error)')
write_stage1(
    s1="The customer states that a payment of £[AMOUNT] was taken from their account on [DATE] which they did not authorise / which was for an incorrect amount. The customer advises the correct amount should have been £[CORRECT AMOUNT] and that the additional £[DIFFERENCE] was taken without authorisation. The customer is requesting the difference to be returned and an explanation of how the error occurred.",
    s2="I have reviewed the customer's account and the payment processed on [DATE]. I reviewed the payment instruction and confirmed that [the direct debit was processed for £[AMOUNT] against an instruction of £[CORRECT AMOUNT] / the standing order was processed for an incorrect amount due to a system update error]. I confirmed this was an error on Santander's side and that the customer did not authorise the additional amount taken.",
    s3="I have confirmed to the customer that the overpayment of £[DIFFERENCE] has been identified as an error and have arranged for the full amount to be returned to their account within [X] working days. I have also reviewed the payment instruction to ensure it is set up correctly going forward and confirmed that future payments will be processed at the correct amount of £[CORRECT AMOUNT].",
    s4="This complaint has been upheld as a Santander error. Our investigation confirmed that the payment of £[AMOUNT] was processed incorrectly due to [specific reason]. The customer did not authorise this amount and it should have been £[CORRECT AMOUNT]. I have returned the difference of £[DIFFERENCE] and offered a goodwill gesture in recognition of the inconvenience caused.",
    s5_text="Refund: £[DIFFERENCE] (overpayment returned)\nGoodwill gesture: £[AMOUNT] — inconvenience of incorrect payment\nTotal: £[TOTAL]",
    letter_type="Upheld Stage 1",
    letter_para1="I'm sorry you've had to contact us about the incorrect payment taken from your account.",
    letter_para2="I've looked into what happened and I can see the issue you've raised was caused by [root cause — e.g. an error in how the direct debit instruction was processed on our side]. I understand this has caused you worry and inconvenience, particularly around your account balance, and for this I'm truly sorry.",
    letter_para3="To fix this for you I've arranged for the overpayment of £[DIFFERENCE] to be returned to your account within [X] working days, and I've confirmed the payment instruction has been corrected so future payments will be taken at the right amount.",
)

divider()
h2('3B — Stage 1 · Declined (customer provided incorrect payee details)')
write_stage1(
    s1="The customer states that a payment of £[AMOUNT] was sent to the wrong account on [DATE] and is requesting that Santander recover the funds. The customer believes the bank is responsible for ensuring the payment was sent to the correct recipient.",
    s2="I have reviewed the customer's account and the payment processed on [DATE]. I reviewed the payment instruction submitted by the customer and confirmed that the sort code [SORT CODE] and account number [ACCOUNT NUMBER] provided by the customer did not match the intended recipient. I confirmed that Santander processed the payment exactly as instructed. I have contacted our payments team and attempted recovery of the funds from the receiving bank. [The receiving bank has confirmed the funds are no longer available as the recipient has spent them / Recovery is ongoing.]",
    s3="I have advised the customer that Santander processed the payment exactly as instructed with the details they provided and that the misdirection was not caused by an error on our part. I have initiated a payment recovery request with the receiving bank on the customer's behalf and will update them on progress. I have provided the customer with the receiving bank's details to contact them directly to request return of the funds.",
    s4="This complaint has not been upheld as there is no Santander error. Our investigation confirmed the payment was processed using the sort code and account number provided by the customer and that these details did not correspond to the intended recipient. Santander is not liable for misdirected payments resulting from incorrect details provided by the customer. We have, as a gesture of goodwill, initiated a recovery request with the receiving bank.",
    s5_text="No compensation offered — the payment was processed correctly using the details the customer provided. Recovery request initiated as a goodwill action. Customer directed to contact receiving bank directly.",
    letter_type="Declined Stage 1",
    letter_para1="I'm sorry you've had to contact us about the payment that was sent to an unintended account.",
    letter_para2="I understand this situation has caused you significant concern, but I've looked into your case and this wasn't because of a mistake we made. This is because the payment was processed using the account details you provided, which did not correspond to your intended recipient.",
    letter_para3="To explain this further, I can confirm that Santander processed your payment of £[AMOUNT] on [DATE] using sort code [SORT CODE] and account number [ACCOUNT NUMBER] — exactly as you instructed. As the payment was sent to the details you provided, we are unable to accept responsibility for the misdirection. However, I have raised a recovery request with the receiving bank on your behalf and will be in touch with an update. I would also recommend contacting the receiving bank, [BANK NAME], directly on [NUMBER] to request they return the funds.",
)

doc.add_page_break()

# ══════════════════════════════════════════════════════════════════════════════
# COMPLAINT TYPE 4 — APP FRAUD NOT REIMBURSED
# ══════════════════════════════════════════════════════════════════════════════
h1('4 — APP Fraud / Scam — Reimbursement Dispute')
note('Authorised Push Payment (APP) fraud. Mandatory reimbursement under PS23/3 effective 7 Oct 2024 for FPS and CHAPS payments. Max reimbursement £85,000. 5 business days to reimburse. Optional excess up to £100. Gross negligence: high bar — vulnerable customers cannot have excess applied.')

h2('4A — Stage 1 · Upheld (reimbursement agreed — bank failed its obligations)')
write_stage1(
    s1="The customer states that on [DATE] they were contacted by [a person claiming to be from / a seller on / an investment platform] and were deceived into authorising a payment of £[AMOUNT] to account [SORT CODE / ACCOUNT NUMBER]. The customer states they did not intend to make a voluntary payment to this recipient and that they were manipulated into doing so. The customer is requesting reimbursement of the £[AMOUNT] lost under the bank's APP fraud reimbursement obligations.",
    s2="I have reviewed the customer's account and the payment of £[AMOUNT] made on [DATE]. I reviewed the Confirmation of Payee (CoP) check result at the time of payment and the fraud typology against PSR PS23/3 guidelines. I confirmed that this is a [purchase / impersonation / investment / romance] scam falling within the mandatory reimbursement regime. I reviewed whether any of the reimbursement exceptions apply: [gross negligence — standard not met / complicit in fraud — no evidence / claim made outside 13-month window — within window / standard excess applies — amount noted]. The payment was made via Faster Payments and is therefore within scope of PS23/3.",
    s3="I have upheld the customer's reimbursement claim and arranged for £[REIMBURSEMENT AMOUNT] to be credited to their account within 5 business days. [A £100 excess has been applied as permitted under PS23/3 / No excess has been applied as the customer is a vulnerable consumer.] I have also reported the fraud to our financial crime team for intelligence sharing purposes.",
    s4="This complaint has been upheld. The payment falls within the mandatory APP fraud reimbursement regime under PSR PS23/3 (effective 7 October 2024). The customer was deceived into authorising the payment and none of the reimbursement exceptions apply. The gross negligence bar has not been met — the customer acted reasonably given the sophistication of the fraud. [An excess of £100 has been applied as permitted under the regulations. / No excess applied — customer is a vulnerable consumer.] Reimbursement of £[AMOUNT] is due within 5 business days.",
    s5_text="Reimbursement: £[AMOUNT] (full loss)\n[Less excess: £100]\nTotal credited to customer: £[NET AMOUNT]\n\nNote: liability for reimbursement is split 50/50 between Santander (sending PSP) and the receiving PSP per PS23/3.",
    letter_type="Upheld Stage 1",
    letter_para1="I'm sorry you've had to contact us about the fraudulent payment made from your account.",
    letter_para2="I've looked into what happened and I can see that you were the victim of a [purchase / impersonation / investment] scam, where you were deceived into authorising a payment to a fraudster. I understand this has caused you significant financial worry and distress, and for this I'm truly sorry.",
    letter_para3="To fix this for you I've arranged for £[NET AMOUNT] to be credited to your account within 5 business days under our APP fraud reimbursement obligations. [A £100 excess has been applied to this payment as permitted under the PSR reimbursement rules.]",
)

divider()
h2('4B — Stage 1 · Declined (gross negligence or outside scope)')
write_stage1(
    s1="The customer states that on [DATE] they authorised a payment of £[AMOUNT] to [ACCOUNT DETAILS] and that they now believe they were defrauded. The customer is requesting reimbursement of the full amount under APP fraud reimbursement rules.",
    s2="I have reviewed the customer's account and the payment of £[AMOUNT] made on [DATE]. I reviewed the circumstances of the payment against PSR PS23/3 guidelines and the fraud typology. [CHOOSE APPLICABLE: I confirmed that the customer was warned at the time of payment via our in-app fraud warning that this type of payment carries a high risk of fraud, and the customer confirmed they wished to proceed. The customer did not contact us to raise concerns at the time of payment despite clear warning prompts. / The payment was made via [BACS / international transfer] which falls outside the FPS and CHAPS scope of PS23/3 mandatory reimbursement. / The claim was submitted [X] months after the payment, outside the 13-month claim window under PS23/3.]",
    s3="I have explained to the customer that their reimbursement claim cannot be upheld on this occasion because [the gross negligence exception applies as the customer proceeded despite clear fraud warnings / the payment type falls outside the mandatory reimbursement regime / the claim is outside the 13-month claim window]. I have advised the customer of their right to refer this matter to the Financial Ombudsman Service.",
    s4="This complaint has not been upheld. [CHOOSE APPLICABLE: The customer was presented with explicit fraud warnings at the time of payment and confirmed they wished to proceed. In proceeding despite clear warnings the customer has failed to take reasonable steps to protect themselves from fraud — the gross negligence bar has been met in this case. / The payment was made via [payment type] which is not within scope of the PS23/3 mandatory reimbursement regime, which applies to FPS and CHAPS only. / The claim was submitted outside the 13-month claim window specified under PS23/3 and cannot be accepted.]",
    s5_text="No reimbursement offered. [Reason: gross negligence / outside scope / outside 13-month window.] Customer advised of right to refer to FOS.",
    letter_type="Declined Stage 1",
    letter_para1="I'm sorry you've had to contact us about the payment you believe was made as a result of fraud.",
    letter_para2="I understand this situation has caused you significant distress, but I've looked into your case and I'm unable to reimburse you on this occasion. This is because [brief reason — e.g. the payment type falls outside the mandatory reimbursement scheme / you proceeded with the payment after being warned of the fraud risk].",
    letter_para3="To explain this further, [full explanation per the investigation findings above]. I understand this will be disappointing to hear. If you disagree with our decision you have the right to refer your complaint to the Financial Ombudsman Service within 6 months of this response.",
)

doc.add_page_break()

# ══════════════════════════════════════════════════════════════════════════════
# COMPLAINT TYPE 5 — MANDATE CHANGE ERROR OR DELAY
# ══════════════════════════════════════════════════════════════════════════════
h1('5 — Mandate Change Error or Delay')
note('Covers: delay in processing a mandate change, incorrect signatories applied, cooling-off period not properly applied, co-signatory not notified.')

h2('5A — Stage 1 · Upheld (Santander delay or error in mandate processing)')
write_stage1(
    s1="The customer states that they submitted a mandate change request on [DATE] to [add / remove / change the signing rule for] [SIGNATORY NAME / ROLE] on their account. They advise that as of [DATE] the change had not been applied and that this has prevented them from completing [business transactions / account operations]. The customer is requesting the mandate change to be applied and compensation for the disruption caused.",
    s2="I have reviewed the customer's account and the mandate change request submitted on [DATE]. I reviewed our internal mandate processing records and confirmed that the request was received on [DATE] and should have been processed within [X] working days. I confirmed that the mandate change had not been applied as of [DATE] due to [an internal processing error / the request being routed incorrectly / required documentation not being requested from the customer at the correct stage]. I confirmed this was an error on Santander's side.",
    s3="I have arranged for the mandate change to be applied as a priority and confirmed to the customer that [SIGNATORY NAME] has been [added / removed / the signing rule has been updated] as of [DATE]. I have also offered a goodwill gesture in recognition of the disruption caused to the customer's business operations during the period the mandate was not updated.",
    s4="This complaint has been upheld as a Santander error. Our investigation confirmed the mandate change was not processed within the expected timeframe due to [specific reason]. This caused disruption to the customer's ability to conduct normal business banking operations. I have resolved the issue and offered a goodwill gesture in recognition of the business inconvenience caused.",
    s5_text="Goodwill gesture: £[AMOUNT] — disruption to business operations caused by our failure to process the mandate change within the expected timeframe.\nTotal: £[AMOUNT]",
    letter_type="Upheld Stage 1",
    letter_para1="I'm sorry you've had to contact us about the delay in updating the mandate on your account.",
    letter_para2="I've looked into what happened and I can see the delay in processing your mandate change was caused by [root cause — e.g. an internal routing error on our side]. I understand this has caused you disruption to your business operations and inconvenience, and for this I'm truly sorry.",
    letter_para3="To fix this for you I've ensured your mandate has now been updated and [SIGNATORY NAME] has been [added / removed] as of today. I've also arranged a goodwill payment of £[AMOUNT] to recognise the disruption this has caused.",
)

divider()
h2('5B — Stage 1 · Declined (customer did not provide required documentation)')
write_stage1(
    s1="The customer states that they submitted a mandate change request on [DATE] and that the change has not been applied. The customer believes the bank has unreasonably delayed the change and is requesting it to be applied immediately.",
    s2="I have reviewed the customer's account and the mandate change request submitted on [DATE]. I reviewed our internal mandate processing records and confirmed that the request was received and that our team wrote to the customer on [DATE] requesting [the required documentation — e.g. updated board minutes / certified ID for the new signatory / all existing signatories' authorisation per the current signing rule]. I confirmed that as of [DATE] the required documentation has not been received, which is why the mandate change has not been applied. Santander cannot apply mandate changes without the required authorisation and documentation for regulatory and fraud prevention reasons.",
    s3="I have explained to the customer that the mandate change cannot be applied until the required documentation is received and have confirmed exactly what is needed: [LIST DOCUMENTS]. I have provided clear guidance on how to submit these and confirmed we will process the change within [X] working days of receipt.",
    s4="This complaint has not been upheld as there is no Santander error. Our investigation confirmed we wrote to the customer requesting the required documentation on [DATE] and this has not yet been received. Santander is required to verify the identity of new signatories and obtain proper authorisation before making changes to account mandates — this is a regulatory obligation, not a discretionary delay.",
    s5_text="No compensation offered — the mandate change has not been applied due to outstanding documentation requested from the customer, not due to a Santander error.",
    letter_type="Declined Stage 1",
    letter_para1="I'm sorry you've had to contact us about the delay in updating the mandate on your account.",
    letter_para2="I understand this situation has caused you frustration, but I've looked into your case and the delay isn't because of a mistake we made. This is because we are still waiting for the required documentation before we can apply the change.",
    letter_para3="To explain this further, our team wrote to you on [DATE] requesting [documents required]. We haven't yet received these, which is why the mandate change hasn't been applied. We have a regulatory obligation to verify the identity of new signatories and obtain proper authorisation before making changes — this protects your business from unauthorised mandate changes. Once we receive [documents], we will apply the change within [X] working days.",
)

doc.add_page_break()

# ══════════════════════════════════════════════════════════════════════════════
# COMPLAINT TYPE 6 — POOR SERVICE / COMMUNICATION
# ══════════════════════════════════════════════════════════════════════════════
h1('6 — Poor Service or Communication')
note('Covers: long wait times, unhelpful or incorrect advice given, failure to follow up, complaint not handled correctly first time, rude or dismissive staff.')

h2('6A — Stage 1 · Upheld (service failure identified)')
write_stage1(
    s1="The customer states that on [DATE] they contacted us regarding [ISSUE] and that the service they received was unsatisfactory. The customer advises that [they were kept waiting for an unreasonable period / they were given incorrect information / they were not followed up as promised / they felt they were treated dismissively]. The customer is requesting an apology and assurance that the issue has been addressed.",
    s2="I have reviewed the customer's account and the interaction on [DATE]. I reviewed [the call recording / the case notes / the correspondence] from [DATE] and confirmed that [the customer waited [X] minutes before being connected / the customer was incorrectly advised that [INCORRECT INFORMATION] / a callback was promised for [DATE] and was not made / the tone of the interaction did not meet our expected service standards]. I confirmed this represents a service failure on our part.",
    s3="I have apologised to the customer for the service failure and confirmed the correct information regarding [ISSUE]. [I have also arranged a goodwill gesture of £[AMOUNT] in recognition of the inconvenience and poor experience caused / I have provided a formal written apology.] I have fed back to the relevant team to ensure this does not happen again.",
    s4="This complaint has been upheld as a Santander service failure. Our review confirmed [specific finding]. The customer did not receive the level of service we expect to deliver and this is not acceptable. [I have offered a goodwill gesture of £[AMOUNT] in recognition of the poor experience caused / The service failure has been acknowledged and a formal apology issued. As the customer has not suffered a direct financial loss, compensation is not appropriate but a goodwill gesture has been offered in recognition of the inconvenience.]",
    s5_text="Goodwill gesture: £[AMOUNT] — recognition of the poor service experience\nOR: No financial compensation — no financial loss suffered; formal written apology issued.",
    letter_type="Upheld Stage 1",
    letter_para1="I'm sorry you've had to contact us about the service you received on [DATE].",
    letter_para2="I've looked into what happened and I can see the poor experience you had was caused by [root cause — e.g. a failure to follow up on our promised callback / incorrect information being provided]. I understand this has caused you frustration and a lack of confidence in how we handle your account, and for this I'm truly sorry.",
    letter_para3="To fix this for you I've [confirmed the correct information regarding [ISSUE] / arranged for a goodwill payment of £[AMOUNT] to your account / fed back to the relevant team]. I can assure you that the experience you described does not reflect the standard of service we aim to provide.",
)

divider()
h2('6B — Stage 1 · Declined (service was within acceptable standards)')
write_stage1(
    s1="The customer states that the service they received on [DATE] when contacting us about [ISSUE] was poor and that they feel their query was not handled appropriately. The customer is requesting compensation for the experience.",
    s2="I have reviewed the customer's account and the interaction on [DATE]. I reviewed [the call recording / the case notes] and confirmed that the customer's query was handled by [TEAM / ADVISOR] and that the information provided was accurate and the service delivered was in line with our expected standards. The wait time of [X] minutes is within our published service levels for [this time of year / this channel]. I did not identify any failure in the service provided.",
    s3="I have explained to the customer that having reviewed the interaction I was unable to identify a service failure and that the information provided to them was correct. I have taken the opportunity to clarify [any point of confusion] and have thanked the customer for bringing their concerns to our attention.",
    s4="This complaint has not been upheld. Our review of the interaction confirmed the service was delivered within our expected standards and the information provided was accurate. As no service failure has been identified we are not in a position to offer compensation on this occasion.",
    s5_text="No compensation offered — no service failure identified.",
    letter_type="Declined Stage 1",
    letter_para1="I'm sorry you've had to contact us about the experience you had when you spoke with us on [DATE].",
    letter_para2="I understand you felt the service didn't meet your expectations, but having listened to the call / reviewed the interaction I wasn't able to identify a service failure on our part. This is because [brief reason — e.g. the wait time was within our published service levels / the information provided was correct].",
    letter_para3="To explain this further, [full explanation]. I do appreciate you taking the time to share your feedback with us and I've made sure it has been noted.",
)

doc.add_page_break()

# ══════════════════════════════════════════════════════════════════════════════
# COMPLAINT TYPE 7 — ACCOUNT CLOSURE DELAY OR ERROR
# ══════════════════════════════════════════════════════════════════════════════
h1('7 — Account Closure Delay or Error')
note('Covers: closure not processed, funds not returned after closure, account still showing charges after closure instruction.')

h2('7A — Stage 1 · Upheld (closure not processed / funds not returned)')
write_stage1(
    s1="The customer states that they submitted a request to close their [ACCOUNT TYPE] account on [DATE] and that as of [DATE] the account has not been closed. The customer advises that [funds of £[AMOUNT] have not been returned to them / charges have continued to be applied since the closure instruction / the account is still visible and active]. The customer is requesting the account to be closed immediately and any outstanding balance returned.",
    s2="I have reviewed the customer's account and the closure request submitted on [DATE]. I reviewed our internal account closure processing records and confirmed that [the closure instruction was received but not actioned due to [REASON] / the account was closed on [DATE] but the remaining balance of £[AMOUNT] was not transferred to the customer's nominated account / charges were applied after the closure instruction was received in error]. This represents an error on Santander's side.",
    s3="I have arranged for the account to be closed immediately and confirmed the outstanding balance of £[AMOUNT] will be transferred to [NOMINATED ACCOUNT] within [X] working days. [I have also arranged a refund of £[AMOUNT] in charges that were applied after the closure instruction was received.] I have offered a goodwill gesture of £[AMOUNT] in recognition of the inconvenience and delay caused.",
    s4="This complaint has been upheld as a Santander error. Our investigation confirmed the account closure was not processed correctly because [specific reason], resulting in [delay / continued charges / funds not being returned]. I have resolved the issue and offered a goodwill gesture in recognition of the inconvenience caused.",
    s5_text="Refund of charges applied after closure instruction: £[AMOUNT]\nGoodwill gesture: £[AMOUNT] — delay in processing closure\nTotal: £[TOTAL]",
    letter_type="Upheld Stage 1",
    letter_para1="I'm sorry you've had to contact us about the delay in closing your account.",
    letter_para2="I've looked into what happened and I can see the delay was caused by [root cause — e.g. an error in how your closure instruction was processed on our side]. I understand this has caused you frustration and inconvenience, and for this I'm truly sorry.",
    letter_para3="To fix this for you I've arranged for your account to be closed and your balance of £[AMOUNT] to be transferred to your nominated account within [X] working days. I've also arranged a goodwill payment of £[AMOUNT] to recognise the inconvenience of the delay.",
)

# ══════════════════════════════════════════════════════════════════════════════
# ESCALATION EXAMPLES — ADDITIONAL TYPES
# ══════════════════════════════════════════════════════════════════════════════
doc.add_page_break()
h1('8 — Escalation: No Change in Original Resolution (Any Complaint Type)')
note('Use this template when the escalation review reaches the same conclusion as the original Stage 1 decision.')

h2('8A — Escalation · No Change (original decision maintained)')
write_escalation(
    e1="The customer disagrees with the original decision to [uphold / decline] their complaint about [BRIEF DESCRIPTION]. The customer [maintains they are owed more compensation / disputes our finding that no Santander error occurred / believes we have not fully considered [NEW POINT]].",
    e2="On escalation I reviewed the original investigation file and all evidence considered at Stage 1. I also reviewed [any new information provided by the customer / our internal records in more detail / correspondence with [THIRD PARTY]]. Having completed this review I reached the same conclusion as the original handler — [the charge was correctly applied / the delay was caused by the previous provider / the level of compensation offered is fair and reasonable]. [I did not identify any new evidence that would change this conclusion.]",
    e3="I have not changed the resolution. The escalation review confirmed the original decision was correct. [The compensation offer of £[AMOUNT] remains fair and reasonable given the circumstances / The decline stands as no Santander error has been identified.] I have explained my findings to the customer and advised them of their right to refer to the FOS.",
    e4_text="No change to original offer of £[AMOUNT] — confirmed as fair and reasonable.\nOR: No compensation offered — decline maintained on escalation review.",
    letter_type="Escalation · No Change",
    letter_para="I'm sorry you don't agree with the resolution previously offered about [brief explanation of what the complaint was about].\n\nI've reviewed your case again in full and have considered [any new information you provided]. Having done so, I've reached the same conclusion as my colleague. [Taking all the circumstances of your complaint into consideration, I believe our original offer of £[AMOUNT] is fair and reasonable. / The delay in [X] happened after [explanation], and our records confirm this was not caused by an error on our part.]\n\nIf you remain unhappy with our decision, you have the right to refer your complaint to the Financial Ombudsman Service within 6 months of this letter.",
)

# ══════════════════════════════════════════════════════════════════════════════
# FOOTER
# ══════════════════════════════════════════════════════════════════════════════
doc.add_page_break()
p = doc.add_paragraph()
shade_paragraph(p, 'C8102E')
p.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = p.add_run('Santander Business Banking · Complaint Writing Guide · Internal Use Only · June 2026')
set_font(run, size=9, color=WHITE)

p = doc.add_paragraph()
p.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.CENTER
p.paragraph_format.space_before = Pt(10)
run = p.add_run('All customer names, account details, dates, and amounts in this guide are illustrative placeholders only. Replace all [BRACKETED] items with real case details before use. Do not include real customer data in training materials.')
set_font(run, size=9, italic=True, color=GREY)

doc.save('Santander_Complaint_Guide.docx')
print('Saved → Santander_Complaint_Guide.docx')
