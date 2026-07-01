#!/usr/bin/env python3
"""Render the Technical Report and Position Paper to .docx from report_content."""
from docx import Document
from docx.shared import Pt, RGBColor, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_TAB_ALIGNMENT
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.enum.section import WD_SECTION
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
import report_content as C

RED = RGBColor(0xEC, 0x00, 0x00); INK = RGBColor(0x1C, 0x19, 0x17)
STONE = RGBColor(0x57, 0x53, 0x4E); LIGHT = RGBColor(0x78, 0x71, 0x6C)
HAIR = 'E7E5E4'; PANEL = 'FBF1EA'


def _shade(el, hexfill):
    tcPr = el.get_or_add_tcPr()
    sh = OxmlElement('w:shd'); sh.set(qn('w:val'), 'clear'); sh.set(qn('w:fill'), hexfill); tcPr.append(sh)

def _border(cell, edges, color=HAIR, sz='6'):
    tcPr = cell._tc.get_or_add_tcPr(); tcB = OxmlElement('w:tcBorders')
    for e in edges:
        x = OxmlElement(f'w:{e}'); x.set(qn('w:val'), 'single'); x.set(qn('w:sz'), sz)
        x.set(qn('w:space'), '0'); x.set(qn('w:color'), color); tcB.append(x)
    tcPr.append(tcB)

def _left_accent(cell, color='EC0000', sz='24'):
    tcPr = cell._tc.get_or_add_tcPr(); tcB = OxmlElement('w:tcBorders')
    x = OxmlElement('w:left'); x.set(qn('w:val'), 'single'); x.set(qn('w:sz'), sz)
    x.set(qn('w:space'), '0'); x.set(qn('w:color'), color); tcB.append(x); tcPr.append(tcB)

def _page_field(run):
    for kind, txt in [('begin', None), (None, 'PAGE'), ('end', None)]:
        if kind:
            fc = OxmlElement('w:fldChar'); fc.set(qn('w:fldCharType'), kind); run._r.append(fc)
        else:
            it = OxmlElement('w:instrText'); it.set(qn('xml:space'), 'preserve'); it.text = txt; run._r.append(it)


def build(meta, body, out):
    doc = Document()
    sec = doc.sections[0]
    sec.top_margin = Inches(0.9); sec.bottom_margin = Inches(0.85)
    sec.left_margin = Inches(0.9); sec.right_margin = Inches(0.9)
    sec.different_first_page_header_footer = True  # cover clean

    st = doc.styles['Normal']; st.font.name = 'Santander Text'; st.font.size = Pt(10.5)
    st.font.color.rgb = INK; st.paragraph_format.space_after = Pt(6); st.paragraph_format.line_spacing = 1.22

    # Running header (title left, classification right)
    hp = sec.header.paragraphs[0]; hp.text = ''
    tabs = hp.paragraph_format.tab_stops
    tabs.add_tab_stop(Inches(6.7), WD_TAB_ALIGNMENT.RIGHT)
    r = hp.add_run(meta['running']); r.font.size = Pt(7.5); r.font.color.rgb = LIGHT
    r2 = hp.add_run('\t' + meta['classification']); r2.font.size = Pt(7.5); r2.font.bold = True; r2.font.color.rgb = RED
    pbd = hp._p.get_or_add_pPr(); b = OxmlElement('w:pBdr'); bo = OxmlElement('w:bottom')
    bo.set(qn('w:val'), 'single'); bo.set(qn('w:sz'), '4'); bo.set(qn('w:space'), '4'); bo.set(qn('w:color'), HAIR)
    b.append(bo); pbd.append(b)

    # Footer (author left, page right)
    fp = sec.footer.paragraphs[0]; fp.text = ''
    fp.paragraph_format.tab_stops.add_tab_stop(Inches(6.7), WD_TAB_ALIGNMENT.RIGHT)
    fr = fp.add_run('Alan Davidson  ·  Internal · Confidential'); fr.font.size = Pt(7.5); fr.font.color.rgb = LIGHT
    fr2 = fp.add_run('\tPage '); fr2.font.size = Pt(7.5); fr2.font.color.rgb = LIGHT
    _page_field(fp.add_run())

    def para(space_before=0, space_after=6, align=None):
        p = doc.add_paragraph(); p.paragraph_format.space_before = Pt(space_before)
        p.paragraph_format.space_after = Pt(space_after)
        if align is not None: p.alignment = align
        return p

    def run(p, text, size=10.5, bold=False, color=INK, italic=False, name='Santander Text'):
        r = p.add_run(text); r.font.size = Pt(size); r.font.bold = bold; r.font.italic = italic
        r.font.color.rgb = color; r.font.name = name; return r

    # ---------- COVER ----------
    doc.add_paragraph()
    logo_p = doc.add_paragraph(); logo_p.paragraph_format.space_after = Pt(4)
    logo_p.add_run().add_picture('brand/santander-logo-red.png', width=Inches(1.9))
    for _ in range(2): doc.add_paragraph()
    run(para(space_after=10), meta['classification'], 9, True, RED)
    run(para(space_after=0), meta['title'], 30, True, INK)
    run(para(space_after=3), meta['subtitle'], 17, False, STONE)
    run(para(space_after=18), meta['tagline'], 11, False, LIGHT, italic=True)
    # document control block
    ctrl = doc.add_table(rows=0, cols=2)
    for k, v in meta['control']:
        cells = ctrl.add_row().cells
        _shade(cells[0]._tc, PANEL); _border(cells[0], ['top', 'bottom'], HAIR); _border(cells[1], ['top', 'bottom'], HAIR)
        run(cells[0].paragraphs[0], k, 8.5, True, STONE)
        run(cells[1].paragraphs[0], v, 8.5, False, INK)
    ctrl.columns[0].cells[0].width = Inches(1.7);
    for c in ctrl.columns[1].cells: c.width = Inches(5.0)
    # contents
    para(space_before=22, space_after=4)
    run(doc.paragraphs[-1], 'Contents', 12, True, INK)
    for item in meta['contents']:
        p = doc.add_paragraph(); p.paragraph_format.space_after = Pt(2)
        run(p, item, 10, False, STONE)
    doc.add_page_break()

    # ---------- BODY ----------
    for blk in body:
        kind = blk[0]
        if kind == 'eyebrow':
            p = para(space_before=14, space_after=1); run(p, blk[1].upper(), 8, True, RED)
        elif kind == 'h1':
            num, title = blk[1], blk[2]
            p = para(space_before=2, space_after=5)
            run(p, (f'{num}   ' if num else '') + title, 15, True, INK)
            pbd = p._p.get_or_add_pPr(); b = OxmlElement('w:pBdr'); bo = OxmlElement('w:bottom')
            bo.set(qn('w:val'), 'single'); bo.set(qn('w:sz'), '6'); bo.set(qn('w:space'), '4'); bo.set(qn('w:color'), HAIR)
            b.append(bo); pbd.append(b)
        elif kind == 'h2':
            p = para(space_before=9, space_after=2); run(p, blk[1], 11.5, True, STONE)
        elif kind == 'lead':
            p = para(space_after=7); run(p, blk[1], 12, False, INK); p.paragraph_format.line_spacing = 1.3
        elif kind == 'p':
            p = para(); run(p, blk[1], 10.5, False, INK)
        elif kind == 'ul':
            for it in blk[1]:
                p = doc.add_paragraph(style='List Bullet'); p.paragraph_format.space_after = Pt(3)
                run(p, it, 10.5, False, INK)
        elif kind == 'callout':
            t = doc.add_table(rows=1, cols=1); t.alignment = WD_TABLE_ALIGNMENT.CENTER
            cell = t.rows[0].cells[0]; _shade(cell._tc, PANEL); _left_accent(cell)
            cell.width = Inches(6.6)
            cp = cell.paragraphs[0]; cp.paragraph_format.space_after = Pt(2)
            run(cp, blk[1].upper() + '  ', 8.5, True, RED)
            cp2 = cell.add_paragraph(); run(cp2, blk[2], 10.5, False, INK); cp2.paragraph_format.space_after = Pt(2)
            doc.add_paragraph().paragraph_format.space_after = Pt(2)
        elif kind == 'table':
            headers, rows, widths = blk[1], blk[2], blk[3]
            t = doc.add_table(rows=1, cols=len(headers)); t.alignment = WD_TABLE_ALIGNMENT.CENTER; t.style = 'Table Grid'
            for i, h in enumerate(t.rows[0].cells):
                _shade(h._tc, '1C1917'); run(h.paragraphs[0], headers[i], 9, True, RGBColor(0xFF, 0xFF, 0xFF))
            for row in rows:
                cs = t.add_row().cells
                for i, v in enumerate(row):
                    run(cs[i].paragraphs[0], v, 9, i == 0, INK)
            for i, w in enumerate(widths):
                for c in t.columns[i].cells: c.width = Inches(w / 25.4)
            doc.add_paragraph().paragraph_format.space_after = Pt(2)
        elif kind == 'refs':
            for n, text in blk[1]:
                p = doc.add_paragraph(); p.paragraph_format.space_after = Pt(3)
                p.paragraph_format.left_indent = Inches(0.35); p.paragraph_format.first_line_indent = Inches(-0.35)
                run(p, f'[{n}]  ', 9, True, STONE); run(p, text, 9, False, INK)
        elif kind == 'spacer':
            doc.add_paragraph()

    # closing note
    doc.add_paragraph()
    p = doc.add_paragraph(); run(p, C.FOOTER_NOTE, 8, False, LIGHT, italic=True)
    doc.save(out); print('WROTE', out)


build(C.TECH_META, C.TECH_BODY, 'Santander_Technical_Report.docx')
build(C.POS_META, C.POS_BODY, 'Santander_Position_Paper.docx')
